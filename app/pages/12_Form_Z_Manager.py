import hashlib
import json
import sys
import time
from pathlib import Path
from typing import Dict, List

import streamlit as st

from app.components.layout import init_page

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    from docx import Document
except ImportError:
    Document = None

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.append(str(ROOT))

from app.components.project_state import (
    ensure_dirs,
    load_manifest,
    project_subdir,
    register_file,
    update_project_meta,
    use_project,
)
from app.components.autofill import d
from app.components.file_utils import build_step_filename, snapshot_to_pdf
from app.components.web3_client import send_log_tx


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# Column mapping for Form Z template (based on BIHi275 example)
# Data starts at row 4, columns 1-14
COLUMNS = [
    "Lfd. Nr.",                                    # Col 1
    "Spender — Bezeichnung (Species)",             # Col 2
    "Spender — RG",                                # Col 3
    "Empfänger — Bezeichnung (Parental_cellline)", # Col 4
    "Empfänger — RG",                              # Col 5
    "Vektor — Bezeichnung (CRISPR-Cas9)",          # Col 6
    "Vektor — Bezeichnung (sgRNA_Sequence)",       # Col 7
    "übertragene Nukleinsäure — Bezeichnung (Gene_Name)",  # Col 8
    "übertragene Nukleinsäure — Bezeichnung (Donor Sequence)",  # Col 9
    "Gefährdungspotential (Hazard potential)",     # Col 10
    "GVO — Bezeichnung (Edited_CellLine)",         # Col 11
    "GVO — RG",                                    # Col 12
    "GVO — erzeugt oder erhalten am",              # Col 13
    "GVO — entsorgt am",                           # Col 14
]


def evaluate_readiness(rows: List[Dict[str, str]]) -> Dict[str, List[str] | bool]:
    issues: List[str] = []
    warnings: List[str] = []
    required = [
        "Lfd. Nr.",
        "Spender — Bezeichnung (Species)",
        "Empfänger — Bezeichnung (Parental_cellline)",
        "Vektor — Bezeichnung (CRISPR-Cas9)",
        "übertragene Nukleinsäure — Bezeichnung (Gene_Name)",
        "GVO — Bezeichnung (Edited_CellLine)",
        "GVO — erzeugt oder erhalten am",
    ]
    for idx, row in enumerate(rows, start=1):
        for field in required:
            if not (row.get(field) or "").strip():
                issues.append(f"Row {idx}: '{field}' is required.")
        if not (row.get("Spender — RG") or "").strip():
            warnings.append(f"Row {idx}: Spender RG missing.")
        if not (row.get("Empfänger — RG") or "").strip():
            warnings.append(f"Row {idx}: Empfänger RG missing.")
        if not (row.get("GVO — RG") or "").strip():
            warnings.append(f"Row {idx}: GVO RG missing.")
    return {"ready": not issues, "issues": issues, "warnings": warnings}

init_page("Step 12 - Form Z Manager")
st.title("Step 12 - Form Z Manager")
st.caption("Maintain Form Z entries, save snapshots, and generate Excel and Word versions.")

project_id = use_project("Project")
if not project_id:
    st.warning("Select a project from the sidebar.")
    st.stop()

manifest = load_manifest(project_id)
meta = manifest.get("meta", {})
compliance = meta.get("compliance", {})

project_root = ensure_dirs(project_id)
snapshots_root = project_root / "snapshots"
snapshot_dir = project_subdir(project_id, "snapshots", "form_z")
reports_dir = project_subdir(project_id, "reports", "form_z")

templates_dir = ROOT / "data" / "templates"
templates_dir.mkdir(parents=True, exist_ok=True)

# Template upload section
st.subheader("📤 Upload New Templates")
with st.expander("Upload templates to library"):
    col_excel, col_word = st.columns(2)
    with col_excel:
        st.caption("Upload Excel template (.xlsx, .xlsm)")
        excel_upload = st.file_uploader("Excel template", type=["xlsx", "xlsm"], key="excel_upload")
        if excel_upload:
            save_path = templates_dir / excel_upload.name
            with save_path.open("wb") as f:
                f.write(excel_upload.getbuffer())
            st.success(f"✅ Saved: {excel_upload.name}")
            st.rerun()
    
    with col_word:
        st.caption("Upload Word template (.docx)")
        word_upload = st.file_uploader("Word template", type=["docx"], key="word_upload")
        if word_upload:
            save_path = templates_dir / word_upload.name
            with save_path.open("wb") as f:
                f.write(word_upload.getbuffer())
            st.success(f"✅ Saved: {word_upload.name}")
            st.rerun()

st.subheader("📋 Select Templates")
available_excel_templates = sorted(list(templates_dir.glob("*.xlsm")) + list(templates_dir.glob("*.xlsx")))
available_word_templates = sorted(templates_dir.glob("*.docx"))

excel_template_saved = Path(compliance.get("form_z_template_path", "")) if compliance.get("form_z_template_path") else None
word_template_saved = Path(compliance.get("formblatt_z_template_path", "")) if compliance.get("formblatt_z_template_path") else None

if available_excel_templates:
    excel_choices = [str(p) for p in available_excel_templates]
    excel_index = 0
    if excel_template_saved and str(excel_template_saved) in excel_choices:
        excel_index = excel_choices.index(str(excel_template_saved))
    excel_template_path = Path(st.selectbox("Form Z Excel template (.xlsm)", excel_choices, index=excel_index))
else:
    excel_template_path = Path(
        st.text_input(
            "Form Z Excel template (.xlsm)",
            str(excel_template_saved or templates_dir / "Z-Form_ BIHi275-A1-A2.xlsx"),
        )
    )

if available_word_templates:
    word_choices = [str(p) for p in available_word_templates]
    word_index = 0
    if word_template_saved and str(word_template_saved) in word_choices:
        word_index = word_choices.index(str(word_template_saved))
    word_template_path = Path(st.selectbox("Formblatt Z Word template (.docx)", word_choices, index=word_index))
else:
    word_template_path = Path(
        st.text_input(
            "Formblatt Z Word template (.docx)",
            str(word_template_saved or templates_dir / "formblatt_z_S1_projekt3.docx"),
        )
    )

start_row = st.number_input(
    "Start row (Excel sheet)",
    min_value=1,
    value=compliance.get("form_z_start_row", 6),
    step=1,
)


def latest_json(folder: Path) -> Dict:
    if not folder.exists():
        return {}
    files = sorted(folder.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    for file in files:
        try:
            return json.loads(file.read_text(encoding="utf-8"))
        except Exception:
            continue
    return {}


# Load data from various project steps
charter = latest_json(snapshots_root / "charter")
design = latest_json(snapshots_root / "design")
master_bank = latest_json(snapshots_root / "master_bank_registry")
lageso_compliance = latest_json(snapshots_root / "lageso_compliance")

# Extract charter information
primary_cell_line = charter.get("cell_line", meta.get("cell_line", ""))
target_gene = charter.get("target_gene", meta.get("target_gene", ""))
gene_symbol = charter.get("gene_symbol", meta.get("gene_symbol", ""))
modification_type = charter.get("modification_type", meta.get("modification_type", ""))
modification_description = charter.get("modification_description", meta.get("modification_description", ""))

# Extract design information
design_guides = design.get("selected_guides") or []
additional_guides = design.get("additional_guides") or []
all_guides = design_guides + additional_guides
sgrna_sequences = ", ".join([g.get("sequence", "") for g in all_guides if g.get("sequence")])
primary_guide = d(design_guides, 0, "sequence", default="")

# Get mutation info from design
mutation_info = design.get("mutation", {})
gene_symbol = mutation_info.get("gene", gene_symbol)

# Get donor information
donors = design.get("donors", [])
donor_sequence = ""
vector_name = ""
if donors:
    donor_sequence = "; ".join([d.get("sequence", "") for d in donors if d.get("sequence")])
    donor_types = [d.get("donor_type", "") for d in donors if d.get("donor_type")]
    if donor_types:
        vector_name = ", ".join(set(donor_types))

# Extract master bank data - get all cell lines registered
mb_entries = master_bank.get("entries", [])
cell_line_names = []
gvo_dates = []
for entry in mb_entries:
    mb_data = entry.get("data", {})
    hpscreg_name = mb_data.get("HPSCReg_Name", "")
    if hpscreg_name:
        cell_line_names.append(hpscreg_name)
    date_field = mb_data.get("Date_Registered") or mb_data.get("date_registered", "")
    if date_field:
        gvo_dates.append(str(date_field))

# Get first entry data for defaults
mb_data = mb_entries[0].get("data", {}) if mb_entries else {}
first_cell_line = cell_line_names[0] if cell_line_names else mb_data.get("HPSCReg_Name", "")
first_date = gvo_dates[0] if gvo_dates else ""

# Extract LAGESO Compliance data
lageso_rows = lageso_compliance.get("rows", [])
lageso_first = lageso_rows[0] if lageso_rows else {}

# Use LAGESO compliance data if available, otherwise use defaults
parental_from_lageso = lageso_first.get("Parental cell line", primary_cell_line)
gene_from_lageso = lageso_first.get("Gene symbol (NCBI ID)", gene_symbol)
donor_from_lageso = lageso_first.get("Donor/Vector (ssODN seq / Addgene cat. number)", "")
modification_desc_from_lageso = lageso_first.get("Modification description", modification_description)
modification_type_lageso = lageso_first.get("Modification Type", modification_type)
rg_from_lageso = lageso_first.get("RG", "1")
edit_intent = mutation_info.get("edit_intent", "")

# Format gene name with Ensembl ID and modification type (like: SORT1 (ENSG00000134243)_Knockout)
gene_name_formatted = f"{gene_from_lageso}_{modification_type_lageso or edit_intent}" if gene_from_lageso else ""

# Format sgRNA sequences with guide IDs and 5'-3' notation 
# (like: PSEN1_A246E_KI_gRNA1: 5'-TGAGCCACTCAGTCCATTCA-3'; PSEN1_A246E_KI_gRNA2: 5'-ACCTCCCTGAATGGACTGAG-3')
guide_seqs_formatted = []
for guide in all_guides:
    guide_id = guide.get("id", "")
    guide_seq = guide.get("sequence", "")
    if guide_seq:
        if guide_id:
            guide_seqs_formatted.append(f"{guide_id}: 5'-{guide_seq}-3'")
        else:
            guide_seqs_formatted.append(f"5'-{guide_seq}-3'")
sgrna_formatted = "; ".join(guide_seqs_formatted) if guide_seqs_formatted else ""

# Build default rows - one per cell line from Master Bank (like BIHi275 example has rows 4 and 5)
default_rows = []
for idx, cell_line_name in enumerate(cell_line_names if cell_line_names else [first_cell_line or project_id]):
    row_date = gvo_dates[idx] if idx < len(gvo_dates) else first_date
    
    default_row = {
        "Lfd. Nr.": str(idx + 1),
        "Spender — Bezeichnung (Species)": "Human",
        "Spender — RG": rg_from_lageso,
        "Empfänger — Bezeichnung (Parental_cellline)": f"Human ({parental_from_lageso})" if parental_from_lageso else "Human",
        "Empfänger — RG": rg_from_lageso,
        "Vektor — Bezeichnung (CRISPR-Cas9)": f"RNP complex (HiFi-Cas9+{gene_from_lageso}_sgRNA)" if gene_from_lageso else "RNP complex",
        "Vektor — Bezeichnung (sgRNA_Sequence)": sgrna_formatted,
        "übertragene Nukleinsäure — Bezeichnung (Gene_Name)": gene_name_formatted,
        "übertragene Nukleinsäure — Bezeichnung (Donor Sequence)": donor_from_lageso or donor_sequence,
        "Gefährdungspotential (Hazard potential)": compliance.get("hazard_potential", "Nein; weil: kein onko- / toxingene"),
        "GVO — Bezeichnung (Edited_CellLine)": cell_line_name,
        "GVO — RG": rg_from_lageso,
        "GVO — erzeugt oder erhalten am": row_date,
        "GVO — entsorgt am": "",
    }
    default_rows.append(default_row)

# Use first row as default if no cell lines
if not default_rows:
    default_rows = [{
        "Lfd. Nr.": "1",
        "Spender — Bezeichnung (Species)": "Human",
        "Spender — RG": rg_from_lageso,
        "Empfänger — Bezeichnung (Parental_cellline)": f"Human ({parental_from_lageso})" if parental_from_lageso else "Human",
        "Empfänger — RG": rg_from_lageso,
        "Vektor — Bezeichnung (CRISPR-Cas9)": "RNP complex",
        "Vektor — Bezeichnung (sgRNA_Sequence)": sgrna_formatted,
        "übertragene Nukleinsäure — Bezeichnung (Gene_Name)": gene_name_formatted,
        "übertragene Nukleinsäure — Bezeichnung (Donor Sequence)": donor_from_lageso or donor_sequence,
        "Gefährdungspotential (Hazard potential)": compliance.get("hazard_potential", "Nein; weil: kein onko- / toxingene"),
        "GVO — Bezeichnung (Edited_CellLine)": first_cell_line or project_id,
        "GVO — RG": rg_from_lageso,
        "GVO — erzeugt oder erhalten am": first_date,
        "GVO — entsorgt am": "",
    }]

# Store auto-filled metadata in session for SD template export
if "form_z_autofill" not in st.session_state:
    st.session_state["form_z_autofill"] = {
        "parental_cell_line": primary_cell_line,
        "cell_line_names": cell_line_names,
        "gene_symbol": gene_symbol or target_gene,
        "sgrna_sequences": sgrna_sequences,
        "donor_sequence": donor_sequence,
        "vector_name": vector_name,
        "modification_type": modification_type,
        "modification_description": modification_description,
        "project_id": project_id,
    }

existing_rows = compliance.get("form_z_rows") or default_rows

st.markdown("### Form Z entries")

# Add refresh button to reload data from latest snapshots
col_refresh, col_info = st.columns([1, 3])
with col_refresh:
    if st.button("🔄 Refresh from latest data", help="Reload table with latest data from Charter, Design, and Master Bank"):
        # Force refresh by clearing old data and using default_rows
        existing_rows = default_rows
        # Update compliance to save the refreshed data
        update_project_meta(project_id, {"compliance": {**compliance, "form_z_rows": default_rows}})
        st.success("✅ Table refreshed with latest project data")
        st.rerun()

with col_info:
    st.caption("💡 Click 'Refresh from latest data' to update the table with current LAGESO Compliance, Charter, Design, and Master Bank information")

# Show auto-fill info
with st.expander("📋 Auto-filled data from project"):
    col1, col2 = st.columns(2)
    with col1:
        st.caption("**From LAGESO Compliance:**")
        st.write(f"- Lfd.Nr.: `{lageso_first.get('Lfd.Nr.', 'N/A')}`")
        st.write(f"- Parental cell line: `{parental_from_lageso or 'N/A'}`")
        st.write(f"- Gene symbol: `{gene_from_lageso or 'N/A'}`")
        st.write(f"- Donor/Vector: `{donor_from_lageso[:50] + '...' if len(donor_from_lageso) > 50 else donor_from_lageso or 'N/A'}`")
        st.write(f"- RG: `{rg_from_lageso or 'N/A'}`")
        st.caption("**From Project Charter:**")
        st.write(f"- Primary cell line: `{primary_cell_line or 'N/A'}`")
        st.write(f"- Modification type: `{modification_type or 'N/A'}`")
    with col2:
        st.caption("**From Design Log:**")
        st.write(f"- sgRNA sequences: `{sgrna_sequences[:50] + '...' if len(sgrna_sequences) > 50 else sgrna_sequences or 'N/A'}`")
        st.write(f"- Vector name: `{vector_name or 'N/A'}`")
        st.caption("**From Master Bank Registry:**")
        if cell_line_names:
            st.write(f"- Cell lines: `{', '.join(cell_line_names)}`")
            st.write(f"- First date: `{first_date or 'N/A'}`")
        else:
            st.write("- No cell lines registered yet")

table_editor = st.data_editor(
    existing_rows,
    column_config={col: st.column_config.TextColumn() for col in COLUMNS},
    num_rows="dynamic",
    use_container_width=True,
    key="form_z_table",
)
table_rows = table_editor.to_dict("records") if hasattr(table_editor, "to_dict") else list(table_editor)

readiness = evaluate_readiness(table_rows or default_rows)


def save_form_z_snapshot(rows: List[Dict[str, str]]):
    timestamp = int(time.time())
    payload = {
        "project_id": project_id,
        "timestamp_unix": timestamp,
        "rows": rows or [default_row],
    }
    payload_bytes = json.dumps(payload, indent=2).encode("utf-8")
    digest = sha256_bytes(payload_bytes)
    outfile = snapshot_dir / f"{project_id}_form_z_{timestamp}_{digest[:12]}.json"
    outfile.write_bytes(payload_bytes)

    st.session_state["form_z_snapshot"] = {
        "digest": digest,
        "outfile": str(outfile),
        "metadata_uri": outfile.resolve().as_uri(),
        "payload": payload,
    }
    st.session_state["form_z_readiness"] = readiness

    register_file(
        project_id,
        "snapshots",
        {
            "step": "form_z",
            "path": str(outfile),
            "digest": digest,
            "timestamp": timestamp,
        },
    )
    pdf_filename = build_step_filename(project_id, "form-z", timestamp)
    pdf_path = reports_dir / pdf_filename
    snapshot_to_pdf(payload, pdf_path, "Form Z Snapshot")
    register_file(
        project_id,
        "reports",
        {
            "step": "form_z",
            "path": str(pdf_path),
            "timestamp": timestamp,
            "digest": sha256_bytes(pdf_path.read_bytes()),
            "type": "pdf",
        },
    )
    st.success("Form Z snapshot saved and PDF generated.")
    with pdf_path.open("rb") as handle:
        st.download_button("Download Form Z snapshot (PDF)", handle, file_name=pdf_path.name)

    primary = rows[0] if rows else default_row
    update_project_meta(
        project_id,
        {
            "compliance": {
                "form_z_rows": rows,
                "form_z_template_path": str(excel_template_path),
                "formblatt_z_template_path": str(word_template_path),
                "form_z_start_row": start_row,
                "form_z_number": primary.get("Lfd. Nr. in Anlage 13/15 Project S2", ""),
                "spender_rg": primary.get("Spender — RG", ""),
                "empfaenger_rg": primary.get("Empfänger — RG", ""),
                "nucleic_acid_present": primary.get("übertragene Nukleinsäure vorhanden?", ""),
                "nucleic_acid_description": primary.get("übertragene Nukleinsäure — Bezeichnung (Description)", ""),
                "hazard_potential": primary.get("Gefährdungspotential (Hazard potential)", ""),
                "gvo_description": primary.get("GVO — Bezeichnung (Description)", ""),
                "gvo_rg": primary.get("GVO — RG", ""),
                "gvo_generated_on": primary.get("GVO — erzeugt oder entsorgt am", ""),
                "gvo_received_on": primary.get("GVO — erhalten am", ""),
            }
        },
    )


def generate_excel(rows: List[Dict[str, str]]):
    if load_workbook is None:
        st.error("openpyxl is required (`pip install openpyxl`).")
        return
    # Resolve a safe Excel template path. Guard against OSError on exists() for bad inputs.
    template_path = excel_template_path
    try:
        exists = False
        try:
            exists = template_path.exists()
        except OSError:
            exists = False
        if not exists:
            fallback = templates_dir / "Z-Form_ BIHi275-A1-A2.xlsx"
            if fallback.exists():
                template_path = fallback
            else:
                # Any .xlsx/.xlsm in templates dir as last resort
                candidates = sorted(list(templates_dir.glob("*.xlsm")) + list(templates_dir.glob("*.xlsx")))
                if candidates:
                    template_path = candidates[0]
                else:
                    st.error("No Excel template found. Please upload a Form Z template in the Templates section.")
                    return
    except Exception as exc:
        st.warning(f"Invalid template path. Falling back to default. Details: {exc}")
        fallback = templates_dir / "Z-Form_ BIHi275-A1-A2.xlsx"
        if fallback.exists():
            template_path = fallback
        else:
            candidates = sorted(list(templates_dir.glob("*.xlsm")) + list(templates_dir.glob("*.xlsx")))
            if candidates:
                template_path = candidates[0]
            else:
                st.error("No Excel template found. Please upload a Form Z template in the Templates section.")
                return

    wb = load_workbook(template_path, keep_vba=True)
    sheet_name = wb.sheetnames[0]
    ws = wb[sheet_name]

    # Get cell line names from Master Bank Registry
    autofill = st.session_state.get("form_z_autofill", {})
    cell_line_names = autofill.get("cell_line_names", [])
    
    # Write all rows to Excel
    for offset, row in enumerate(rows or [default_row]):
        for col_idx, header in enumerate(COLUMNS, start=1):
            ws.cell(row=start_row + offset, column=col_idx, value=row.get(header, ""))

    timestamp = int(time.time())
    
    # Create filename with all cell line names: Z-Form_ BIHi275-A1-A2.xlsx
    # Extract just the cell line identifiers (like BIHi005-A-1X) and combine with hyphens
    def _safe_filename_component(s: str) -> str:
        s = (s or "").strip()
        # Allow letters, numbers, dash and underscore; replace others with '-'
        s = s.replace(" ", "-").replace("/", "-").replace("\\", "-")
        # Collapse multiple dashes
        while "--" in s:
            s = s.replace("--", "-")
        # Cap component length to 40 chars to keep overall filename reasonable
        return s[:40]

    if cell_line_names:
        formatted_names = [_safe_filename_component(name) for name in cell_line_names]
        combined_name = "-".join(formatted_names)
        # Final safety cap: ensure entire filename (without path) stays < 120 chars
        base_name = f"Z-Form_ {combined_name}.xlsx"
        if len(base_name) > 120:
            # Trim the combined part conservatively
            keep_len = max(1, 120 - len("Z-Form_ .xlsx"))
            combined_name = combined_name[:keep_len]
            base_name = f"Z-Form_ {combined_name}.xlsx"
        out_path = reports_dir / base_name
    else:
        # Fallback if no cell lines found
        out_path = reports_dir / f"{project_id}_FormZ_{timestamp}.xlsx"
        out_path = reports_dir / f"{project_id}_FormZ_{timestamp}.xlsx"
    
    wb.save(out_path)
    wb.close()
    
    register_file(
        project_id,
        "reports",
        {
            "step": "form_z",
            "path": str(out_path),
            "timestamp": timestamp,
            "type": "xlsx",
            "label": "Form Z Excel",
            "cell_lines": cell_line_names,
        },
    )
    st.success(f"✅ Form Z Excel generated: {out_path.name}")
    with out_path.open("rb") as handle:
        st.download_button("Download Form Z Excel", handle, file_name=out_path.name, key=f"formz_excel_{timestamp}")


def generate_word(rows: List[Dict[str, str]]):
    if Document is None:
        st.error("python-docx is required (`pip install python-docx`).")
        return
    if not word_template_path.exists():
        st.error(f"Word template not found: {word_template_path}")
        return

    doc = Document(word_template_path)
    replacements = {
        "{{PROJECT_ID}}": project_id,
        "{{LFD_NR}}": rows[0].get("Lfd. Nr. in Anlage 13/15 Project S2", ""),
        "{{SPENDER}}": rows[0].get("Spender — Bezeichnung (Description)", ""),
        "{{SPENDER_RG}}": rows[0].get("Spender — RG", ""),
        "{{EMPFAENGER}}": rows[0].get("Empfänger — Bezeichnung (Description)", ""),
        "{{EMPFAENGER_RG}}": rows[0].get("Empfänger — RG", ""),
        "{{VECTOR}}": rows[0].get("Vektor — Bezeichnung (Description)", ""),
        "{{NA_VORHANDEN}}": rows[0].get("übertragene Nukleinsäure vorhanden?", ""),
        "{{NA_BESCHREIBUNG}}": rows[0].get("übertragene Nukleinsäure — Bezeichnung (Description)", ""),
        "{{HAZARD}}": rows[0].get("Gefährdungspotential (Hazard potential)", ""),
        "{{GVO}}": rows[0].get("GVO — Bezeichnung (Description)", ""),
        "{{GVO_RG}}": rows[0].get("GVO — RG", ""),
        "{{GVO_GENERIERT}}": rows[0].get("GVO — erzeugt oder entsorgt am", ""),
        "{{GVO_ERHALTEN}}": rows[0].get("GVO — erhalten am", ""),
    }

    def replace(node):
        for paragraph in node.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(key, value)
        for table in node.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace(cell)

    replace(doc)
    timestamp = int(time.time())
    out_path = reports_dir / f"{project_id}_FormblattZ_{timestamp}.docx"
    doc.save(out_path)
    register_file(
        project_id,
        "reports",
        {
            "step": "form_z",
            "path": str(out_path),
            "timestamp": timestamp,
            "type": "docx",
            "label": "Formblatt Z Word",
        },
    )
    st.success(f"Formblatt Z generated: {out_path.name}")
    with out_path.open("rb") as handle:
        st.download_button("Download Formblatt Z", handle, file_name=out_path.name)


def append_to_sd_template(rows: List[Dict[str, str]]):
    """Append project data to Sheet 3 of 13_15 SD template."""
    template_path = templates_dir / "13_15 SD Proj 1-3.xlsm"
    if not template_path.exists():
        st.error(f"Template not found: {template_path}")
        return
    
    # Load template (not read-only so we can write)
    wb = load_workbook(template_path)
    if len(wb.sheetnames) < 3:
        st.error("Template must have at least 3 sheets")
        wb.close()
        return
    
    ws = wb[wb.sheetnames[2]]  # Sheet 3 (0-indexed as sheet 2)
    
    # Find next empty row (skip header at row 3)
    next_row = ws.max_row + 1
    
    # Get auto-filled metadata
    autofill = st.session_state.get("form_z_autofill", {})
    parental_cell_line = autofill.get("parental_cell_line", "")
    cell_line_names = autofill.get("cell_line_names", [])
    gene_symbol = autofill.get("gene_symbol", "")
    sgrna_sequences = autofill.get("sgrna_sequences", "")
    donor_sequence = autofill.get("donor_sequence", "")
    vector_name = autofill.get("vector_name", "")
    modification_type = autofill.get("modification_type", "")
    modification_description = autofill.get("modification_description", "")
    project_id_val = autofill.get("project_id", project_id)
    
    # SD Template columns (Row 3):
    # A: Lfd.Nr., B: Parental cell line, C: Cell line name, D: Link Form Z, 
    # E: IRIS ID, F: Gene symbol, G: crRNA/sgRNA, H: Donor/Vector, I: (empty),
    # J: Modification Type, K: Modification description, L: Zygosity,
    # M: Date of Registration, N: RG, O: Disease
    
    # If multiple cell lines from master bank, create a row for each
    if cell_line_names:
        for idx, cell_line in enumerate(cell_line_names, start=1):
            row_data = rows[idx - 1] if idx - 1 < len(rows) else rows[0] if rows else {}
            
            ws.cell(next_row, 1, idx)  # Lfd.Nr
            ws.cell(next_row, 2, parental_cell_line)  # Parental cell line
            ws.cell(next_row, 3, cell_line)  # Cell line name from master bank
            ws.cell(next_row, 4, "")  # Link Form Z
            ws.cell(next_row, 5, project_id_val)  # IRIS ID (Project ID)
            ws.cell(next_row, 6, gene_symbol)  # Gene symbol from charter
            ws.cell(next_row, 7, sgrna_sequences)  # crRNA/sgRNA from design
            ws.cell(next_row, 8, vector_name or donor_sequence[:100] if donor_sequence else "")  # Donor/Vector
            ws.cell(next_row, 10, modification_type)  # Modification Type from charter
            ws.cell(next_row, 11, modification_description)  # Modification description from charter
            ws.cell(next_row, 12, row_data.get("Zygosity", ""))  # Zygosity (manual entry)
            ws.cell(next_row, 13, row_data.get("GVO — erzeugt oder entsorgt am", ""))  # Date
            ws.cell(next_row, 14, row_data.get("GVO — RG", "1"))  # RG
            ws.cell(next_row, 15, meta.get("disease", ""))  # Disease from project meta
            next_row += 1
    else:
        # No master bank entries, use Form Z data
        for idx, row in enumerate(rows, start=1):
            ws.cell(next_row, 1, idx)  # Lfd.Nr
            ws.cell(next_row, 2, parental_cell_line)  # Parental cell line
            ws.cell(next_row, 3, row.get("GVO — Bezeichnung (Description)", ""))  # Cell line name
            ws.cell(next_row, 4, "")  # Link Form Z
            ws.cell(next_row, 5, project_id_val)  # IRIS ID
            ws.cell(next_row, 6, gene_symbol)  # Gene symbol
            ws.cell(next_row, 7, sgrna_sequences)  # crRNA/sgRNA
            ws.cell(next_row, 8, vector_name or donor_sequence[:100] if donor_sequence else "")  # Donor/Vector
            ws.cell(next_row, 10, modification_type)  # Modification Type
            ws.cell(next_row, 11, modification_description)  # Modification description
            ws.cell(next_row, 12, "")  # Zygosity
            ws.cell(next_row, 13, row.get("GVO — erzeugt oder entsorgt am", ""))  # Date
            ws.cell(next_row, 14, row.get("GVO — RG", "1"))  # RG
            ws.cell(next_row, 15, meta.get("disease", ""))  # Disease
            next_row += 1
    
    # Save to reports directory
    timestamp = int(time.time())
    output_path = reports_dir / f"13_15_SD_updated_{timestamp}.xlsm"
    wb.save(output_path)
    wb.close()
    
    digest = sha256_bytes(output_path.read_bytes())
    register_file(
        project_id,
        "reports",
        {
            "step": "form_z",
            "filename": output_path.name,
            "path": str(output_path),
            "sha256": digest,
            "timestamp": timestamp,
            "type": "xlsm",
            "label": "13/15 SD Template Export",
        },
    )
    st.success(f"✅ Data appended to SD template: {output_path.name}")
    with output_path.open("rb") as handle:
        st.download_button("Download 13/15 SD Template", handle, file_name=output_path.name, key="sd_download")


col_snapshot, col_excel, col_word, col_sd = st.columns(4)
if col_snapshot.button("Save Form Z snapshot"):
    save_form_z_snapshot(table_rows)

if col_excel.button("Generate Excel"):
    generate_excel(table_rows)

if col_word.button("Generate Word"):
    generate_word(table_rows)

if col_sd.button("Export to 13/15 SD"):
    append_to_sd_template(table_rows)

col_anchor = st.container()
snapshot_state = st.session_state.get("form_z_snapshot")
readiness_state = st.session_state.get("form_z_readiness", readiness)
with col_anchor:
    st.markdown("### Anchoring")
    if snapshot_state:
        st.code(f"Snapshot hash: {snapshot_state['digest']}", language="text")
        if readiness_state["ready"]:
            st.success("Ready to anchor.")
        else:
            st.error("Resolve issues before anchoring:")
            for issue in readiness_state["issues"]:
                st.write(f"- {issue}")
            if readiness_state["warnings"]:
                st.caption("Warnings:")
                for warning in readiness_state["warnings"]:
                    st.caption(f"Warning: {warning}")
        disabled = not readiness_state["ready"]
        if st.button("Anchor Form Z snapshot", disabled=disabled):
            try:
                result = send_log_tx(
                    hex_digest=snapshot_state["digest"],
                    step="Form Z",
                    metadata_uri=snapshot_state["metadata_uri"],
                )
            except Exception as exc:
                st.error(f"Anchoring failed: {exc}")
            else:
                st.success("Anchored on-chain.")
                st.write(f"Tx: {result['tx_hash']}")
                st.json(result["receipt"])
                register_chain_tx(
                    project_id,
                    {
                        "step": "form_z",
                        "tx_hash": result["tx_hash"],
                        "digest": snapshot_state["digest"],
                        "timestamp": int(time.time()),
                        "metadata_uri": snapshot_state["metadata_uri"],
                    },
                )
    else:
        st.info("Save a Form Z snapshot before anchoring.")

