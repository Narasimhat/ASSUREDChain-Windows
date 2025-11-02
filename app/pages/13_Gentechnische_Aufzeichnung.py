import hashlib
import json
import sys
import time
from pathlib import Path
from typing import Dict, List

import streamlit as st

from app.components.layout import init_page

try:
    from docx import Document
except ImportError:
    Document = None

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.append(str(ROOT))

from app.components.project_state import (
    append_audit,
    ensure_dirs,
    load_manifest,
    project_subdir,
    register_file,
    register_chain_tx,
    update_project_meta,
    use_project,
)
from app.components.autofill import d
from app.components.file_utils import build_step_filename, snapshot_to_pdf
from app.components.web3_client import send_log_tx


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def latest_json(folder: Path) -> Dict:
    if not folder.exists():
        return {}
    files = sorted(folder.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    for file in files:
        try:
            return json.loads(file.read_text(encoding="utf-8"))
        except Exception:
            continue
    return {}


def evaluate_readiness(payload: Dict) -> Dict[str, List[str] | bool]:
    issues: List[str] = []
    warnings: List[str] = []
    required_fields = [
        "operator_address",
        "facility_location",
        "facility_number",
        "project_lead",
        "biosafety_officer",
        "work_number",
        "work_topic",
        "security_level",
        "permit_date",
        "start_date",
    ]
    for field in required_fields:
        if not (payload.get(field) or "").strip():
            issues.append(f"{field.replace('_', ' ').title()} is required.")

    gvo_rows = payload.get("gvo_rows") or []
    if not gvo_rows:
        issues.append("At least one GVO row must be documented.")
    else:
        for idx, row in enumerate(gvo_rows, start=1):
            if not (row.get("gRNA Description") or "").strip():
                issues.append(f"GVO row {idx}: gRNA description required.")
            if not (row.get("GVO Description") or "").strip():
                issues.append(f"GVO row {idx}: GVO description required.")
            if not (row.get("Generated / Received on") or "").strip():
                warnings.append(f"GVO row {idx}: generation or receipt date missing.")
    return {"ready": not issues, "issues": issues, "warnings": warnings}


def build_document(payload: Dict, out_path: Path) -> Path:
    doc = Document()
    doc.add_heading("AUFZEICHNUNG FÜR EINE GENTECHNISCHE ARBEIT NACH GENTAUFZV", level=1)

    doc.add_paragraph("1. Name und Anschrift des Betreibers")
    doc.add_paragraph(payload.get("operator_address", ""))

    doc.add_paragraph("2. Lage der gentechnischen Anlage")
    doc.add_paragraph(payload.get("facility_location", ""))

    doc.add_paragraph(f"3. Nr. der Anlage: {payload.get('facility_number', '')}")

    doc.add_paragraph("4. Projektleiter (ggf. weitere PL)")
    doc.add_paragraph(payload.get("project_lead", ""))
    for entry in payload.get("additional_project_leads", []):
        doc.add_paragraph(entry, style="List Bullet")

    doc.add_paragraph("5. Beauftragter für die Biologische Sicherheit")
    doc.add_paragraph(payload.get("biosafety_officer", ""))

    doc.add_paragraph(
        "6. Ab Sicherheitsstufe 2: Bei Umgang mit humanpathogenen Organismen Personen,"
        " die in der gentechnischen Anlage tätig sind"
    )
    for person in payload.get("s2_personnel", []):
        doc.add_paragraph(person, style="List Bullet")

    doc.add_paragraph(f"7. Nr. der Arbeit: {payload.get('work_number', '')}")
    doc.add_paragraph("8. Thema der Arbeit")
    doc.add_paragraph(payload.get("work_topic", ""))

    doc.add_paragraph("9. Sicherheitsstufe")
    doc.add_paragraph(payload.get("security_level", ""))

    doc.add_paragraph("10. Datum des Bescheides oder der Eingangsbestätigung")
    doc.add_paragraph(payload.get("permit_date", ""))

    doc.add_paragraph("11. Zeitpunkt des Beginns und Abschlusses der gentechnischen Arbeiten")
    doc.add_paragraph(f"Beginn: {payload.get('start_date', '')}")
    doc.add_paragraph(f"Abschluss: {payload.get('end_date', '')}")

    doc.add_paragraph("12. Besondere Vorkommnisse")
    doc.add_paragraph(payload.get("special_incidents", ""))

    doc.add_paragraph(
        f"13. Angaben zu den gentechnisch veränderten Organismen (GVO) der Arbeit Nr.: {payload.get('work_number', '')} "
        f"in der Anlage Nr.: {payload.get('facility_number', '')}"
    )
    table = doc.add_table(rows=1, cols=8)
    headers = [
        "Lfd. Nr.",
        "gRNA (Beschreibung / Sequenz)",
        "ssODN (Beschreibung / Sequenz)",
        "GVO (Bezeichnung)",
        "RG",
        "Erzeugt / erhalten am",
        "Entsorgt am",
        "Bemerkung",
    ]
    for idx, head in enumerate(headers):
        table.rows[0].cells[idx].text = head
    for row in payload.get("gvo_rows", []):
        cells = table.add_row().cells
        cells[0].text = row.get("Lfd. Nr.", "")
        cells[1].text = f"{row.get('gRNA Description', '')}\n{row.get('gRNA Sequence', '')}"
        cells[2].text = f"{row.get('ssODN Description', '')}\n{row.get('ssODN Sequence', '')}"
        cells[3].text = row.get("GVO Description", "")
        cells[4].text = row.get("GVO RG", "")
        cells[5].text = row.get("Generated / Received on", "")
        cells[6].text = row.get("Disposed on", "")
        cells[7].text = row.get("Notes", "")

    doc.add_paragraph(
        "14. Inaktivierung des Abfalls durch: "
        + payload.get("waste_method", "Autoklavieren")
    )

    doc.add_paragraph("15. Kenntnisnahme")
    doc.add_paragraph(f"Funktion: {payload.get('acknowledgement_role', '')}")
    doc.add_paragraph(f"Datum: {payload.get('acknowledgement_date', '')}")

    doc.add_page_break()
    doc.add_paragraph(
        f"Anlage zu den Aufzeichnungen Arbeit Nr.: {payload.get('work_number', '')} "
        f"Anlage Nr.: {payload.get('facility_number', '')}"
    )
    doc.add_paragraph("Verzeichnis der verwendeten Abkürzungen:")
    abk_table = doc.add_table(rows=1, cols=2)
    abk_table.rows[0].cells[0].text = "Abkürzung"
    abk_table.rows[0].cells[1].text = "Erläuterung"
    for row in payload.get("abbreviations", []):
        cells = abk_table.add_row().cells
        cells[0].text = row.get("Abkürzung", "")
        cells[1].text = row.get("Erläuterung", "")

    doc.add_page_break()
    doc.add_paragraph(
        f"Anlage zu den Aufzeichnungen Arbeit Nr.: {payload.get('work_number', '')} "
        f"Anlage Nr.: {payload.get('facility_number', '')}"
    )
    doc.add_paragraph(
        "Beschreibung der weiteren Arbeiten der Sicherheitsstufe 1 einschließlich Zielsetzung und Risikobewertung:"
    )
    doc.add_paragraph(payload.get("additional_description", ""))

    doc.save(out_path)
    return out_path

init_page("Step 13 - GenTAufzV Work Record")
st.title("Step 13 - GenTAufzV Work Record")
st.caption("Capture the GenTAufzV work log, generate a Word document, snapshot, and optional chain anchor.")

project_id = use_project("Project")
if not project_id:
    st.warning("Select a project from the sidebar.")
    st.stop()

if Document is None:
    st.error("python-docx required. Install with `pip install python-docx`.")
    st.stop()

manifest = load_manifest(project_id)
meta = manifest.get("meta", {})
compliance = meta.get("compliance", {})

project_root = ensure_dirs(project_id)
snapshots_root = project_root / "snapshots"
reports_dir = project_subdir(project_id, "reports", "lageso")
snapshot_dir = project_subdir(project_id, "snapshots", "gentaaufzv")

design = latest_json(snapshots_root / "design")
master_bank = latest_json(snapshots_root / "master_bank_registry")

design_guides = design.get("selected_guides") or []
primary_guide = d(design_guides, 0, "sequence", default="")
donor_sequence = (design.get("donor") or {}).get("sequence", "")

mb_entries = master_bank.get("entries") or []
mb_data = mb_entries[0]["data"] if mb_entries else {}

defaults = compliance.get("aufzeichnung", {})

operator_address = st.text_area(
    "1. Name und Anschrift des Betreibers",
    value=defaults.get("operator_address", "Max Delbrück Centrum für Molekulare Medizin (MDC)"),
)
facility_location = st.text_area(
    "2. Lage der gentechnischen Anlage",
    value=defaults.get("facility_location", ""),
)
facility_number = st.text_input("3. Nr. der Anlage", value=defaults.get("facility_number", "13/15"))
project_lead = st.text_input("4. Projektleiter", value=defaults.get("project_lead", compliance.get("responsible_person", "")))
additional_leads = st.text_area(
    "Weitere Projektleiter (optional, eine Zeile pro Person)",
    value="\n".join(defaults.get("additional_project_leads", [])),
)
biosafety_officer = st.text_input(
    "5. Beauftragter für die Biologische Sicherheit",
    value=defaults.get("biosafety_officer", "Sandra Schommer"),
)
s2_personnel_text = st.text_area(
    "6. (optional) Personen für Sicherheitsstufe 2 (eine Zeile pro Person)",
    value="\n".join(defaults.get("s2_personnel", [])),
)
work_number = st.text_input("7. Nr. der Arbeit", value=defaults.get("work_number", "3"))
work_topic = st.text_area(
    "8. Thema der Arbeit",
    value=defaults.get(
        "work_topic",
        "Genomeditierung von induzierten pluripotenten Zelllinien sowie Herstellung von Reporterkonstrukten.",
    ),
)
security_level = st.text_input("9. Sicherheitsstufe", value=defaults.get("security_level", "S1"))
permit_date = st.text_input("10. Datum des Bescheides", value=defaults.get("permit_date", "12.01.2015"))
start_date = st.text_input("11. Beginn der Arbeiten", value=defaults.get("start_date", "01.06.2018"))
end_date = st.text_input("11. Abschluss der Arbeiten", value=defaults.get("end_date", ""))
special_incidents = st.text_area("12. Besondere Vorkommnisse", value=defaults.get("special_incidents", ""))

st.markdown("### GVO Tabelle")
gvo_defaults = defaults.get("gvo_rows") or [
    {
        "Lfd. Nr.": "1",
        "gRNA Description": "gRNA",
        "gRNA Sequence": primary_guide,
        "ssODN Description": "ssODN",
        "ssODN Sequence": donor_sequence,
        "GVO Description": mb_data.get("IRIS_ID", ""),
        "GVO RG": defaults.get("gvo_rg", ""),
        "Generated / Received on": defaults.get("gvo_generated_on", ""),
        "Disposed on": defaults.get("gvo_disposed_on", ""),
        "Notes": "",
    }
]
gvo_table = st.data_editor(
    gvo_defaults,
    column_config={col: st.column_config.TextColumn() for col in gvo_defaults[0].keys()},
    num_rows="dynamic",
    use_container_width=True,
    key="gvo_table_editor",
)
gvo_rows = gvo_table.to_dict("records") if hasattr(gvo_table, "to_dict") else list(gvo_table)

st.markdown("### Abkürzungsverzeichnis")
abbrev_defaults = defaults.get("abbreviations") or [
    {"Abkürzung": "ATG", "Erläuterung": "Start codon targeting"},
    {"Abkürzung": "GC", "Erläuterung": "Gene correction"},
]
abbrev_table = st.data_editor(
    abbrev_defaults,
    column_config={"Abkürzung": st.column_config.TextColumn(), "Erläuterung": st.column_config.TextColumn()},
    num_rows="dynamic",
    use_container_width=True,
    key="abbrev_table_editor",
)
abbreviations = abbrev_table.to_dict("records") if hasattr(abbrev_table, "to_dict") else list(abbrev_table)

additional_description = st.text_area(
    "Beschreibung der weiteren Arbeiten (Anlagebeschreibung)",
    value=defaults.get("additional_description", ""),
)
waste_method = st.text_input("14. Inaktivierung des Abfalls durch", value=defaults.get("waste_method", "Autoklavieren"))
ack_role = st.text_input("15. Funktion (Kenntnisnahme)", value=defaults.get("acknowledgement_role", ""))
ack_date = st.text_input("15. Datum (Kenntnisnahme)", value=defaults.get("acknowledgement_date", ""))

payload = {
    "operator_address": operator_address,
    "facility_location": facility_location,
    "facility_number": facility_number,
    "project_lead": project_lead,
    "additional_project_leads": [line for line in additional_leads.splitlines() if line.strip()],
    "biosafety_officer": biosafety_officer,
    "s2_personnel": [line for line in s2_personnel_text.splitlines() if line.strip()],
    "work_number": work_number,
    "work_topic": work_topic,
    "security_level": security_level,
    "permit_date": permit_date,
    "start_date": start_date,
    "end_date": end_date,
    "special_incidents": special_incidents,
    "gvo_rows": gvo_rows,
    "abbreviations": abbreviations,
    "additional_description": additional_description,
    "waste_method": waste_method,
    "acknowledgement_role": ack_role,
    "acknowledgement_date": ack_date,
}

readiness = evaluate_readiness(payload)

col_snapshot, col_docx, col_anchor = st.columns(3)

with col_snapshot:
    if st.button("Save work log snapshot"):
        timestamp = int(time.time())
        payload["timestamp_unix"] = timestamp
        payload_bytes = json.dumps(payload, indent=2).encode("utf-8")
        digest = sha256_bytes(payload_bytes)
        outfile = snapshot_dir / f"{project_id}_gentaaufzv_{timestamp}_{digest[:12]}.json"
        outfile.write_bytes(payload_bytes)

        st.session_state["genta_snapshot"] = {
            "digest": digest,
            "outfile": str(outfile),
            "metadata_uri": outfile.resolve().as_uri(),
            "payload": payload,
        }
        st.session_state["genta_readiness"] = readiness

        update_project_meta(
            project_id,
            {
                "compliance": {
                    "aufzeichnung": {
                        **defaults,
                        "operator_address": operator_address,
                        "facility_location": facility_location,
                        "facility_number": facility_number,
                        "project_lead": project_lead,
                        "additional_project_leads": [line for line in additional_leads.splitlines() if line.strip()],
                        "biosafety_officer": biosafety_officer,
                        "s2_personnel": [line for line in s2_personnel_text.splitlines() if line.strip()],
                        "work_number": work_number,
                        "work_topic": work_topic,
                        "security_level": security_level,
                        "permit_date": permit_date,
                        "start_date": start_date,
                        "end_date": end_date,
                        "special_incidents": special_incidents,
                        "gvo_rows": gvo_rows,
                        "abbreviations": abbreviations,
                        "additional_description": additional_description,
                        "waste_method": waste_method,
                        "acknowledgement_role": ack_role,
                        "acknowledgement_date": ack_date,
                    }
                }
            },
        )

        register_file(
            project_id,
            "snapshots",
            {
                "step": "gentaaufzv",
                "path": str(outfile),
                "digest": digest,
                "timestamp": timestamp,
            },
        )

        pdf_filename = build_step_filename(project_id, "gentaaufzv", timestamp)
        pdf_path = reports_dir / pdf_filename
        snapshot_to_pdf(payload, pdf_path, "GenTAufzV Work Log")
        register_file(
            project_id,
            "reports",
            {
                "step": "gentaaufzv",
                "path": str(pdf_path),
                "timestamp": timestamp,
                "digest": sha256_bytes(pdf_path.read_bytes()),
                "type": "pdf",
            },
        )
        st.success("Snapshot and PDF saved.")
        with pdf_path.open("rb") as handle:
            st.download_button("Download PDF", handle, file_name=pdf_path.name)

with col_docx:
    if st.button("Generate Word document"):
        if Document is None:
            st.error("python-docx not available.")
        else:
            timestamp = int(time.time())
            out_path = reports_dir / f"{project_id}_GenTAufzV_{timestamp}.docx"
            try:
                build_document(payload, out_path)
            except Exception as exc:
                st.error(f"Failed to build document: {exc}")
            else:
                register_file(
                    project_id,
                    "reports",
                    {
                        "step": "gentaaufzv",
                        "path": str(out_path),
                        "timestamp": timestamp,
                        "type": "docx",
                        "sha256": sha256_bytes(out_path.read_bytes()),
                    },
                )
                st.success(f"Word file generated: {out_path.name}")
                with out_path.open("rb") as handle:
                    st.download_button("Download DOCX", handle, file_name=out_path.name)

with col_anchor:
    snapshot_state = st.session_state.get("genta_snapshot")
    readiness_state = st.session_state.get("genta_readiness", readiness)
    if snapshot_state:
        st.code(f"Hash: {snapshot_state['digest']}", language="text")
        if readiness_state["ready"]:
            st.success("Ready to anchor.")
        else:
            st.error("Resolve issues before anchoring:")
            for issue in readiness_state["issues"]:
                st.write(f"- {issue}")
            if readiness_state["warnings"]:
                st.caption("Warnings:")
                for warning in readiness_state["warnings"]:
                    st.caption(f"?? {warning}")
        disabled = not readiness_state["ready"]
        if st.button("Anchor snapshot", disabled=disabled):
            try:
                result = send_log_tx(
                    hex_digest=snapshot_state["digest"],
                    step="GenTAufzV Work Log",
                    metadata_uri=snapshot_state["metadata_uri"],
                )
            except Exception as exc:
                st.error(f"Anchoring failed: {exc}")
            else:
                st.success("Anchored on-chain.")
                st.write(f"Tx: {result['tx_hash']}")
                st.json(result['receipt'])
                register_chain_tx(
                    project_id,
                    {
                        "step": "gentaaufzv",
                        "tx_hash": result["tx_hash"],
                        "digest": snapshot_state["digest"],
                        "timestamp": int(time.time()),
                        "metadata_uri": snapshot_state["metadata_uri"],
                    },
                )
                append_audit(
                    project_id,
                    {
                        "timestamp": int(time.time()),
                        "event": "gentaaufzv_anchored",
                        "tx_hash": result["tx_hash"],
                        "file": snapshot_state["outfile"],
                    },
                )
    else:
        st.info("Save a snapshot before anchoring.")

