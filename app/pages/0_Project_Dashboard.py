import hashlib
import io
import json
import time
import zipfile
from pathlib import Path
from typing import Dict, List, Optional

import streamlit as st
from docx import Document
try:
    from PyPDF2 import PdfReader  # for validating PDFs before merge
except Exception:  # pragma: no cover
    PdfReader = None

from app.components.file_utils import build_step_filename, merge_pdfs
from app.components.layout import init_page

from app.components.project_state import (
    append_audit,
    ensure_dirs,
    list_projects,
    load_manifest,
    manifest_path,
    project_subdir,
    register_chain_tx,
    register_file,
    use_project,
)
from app.components.web3_client import send_log_tx
from scripts.regenerate_binder import build_binder


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


ROOT = Path(__file__).resolve().parents[2]
ASSURED_STEP_ORDER: List[tuple[str, str]] = [
    ("charter", "Project Charter"),
    ("design", "Design"),
    ("delivery", "Delivery"),
    ("assessment", "Assessment"),
    ("cloning", "Cloning"),
    # Removed obsolete steps (screening, form_z, gentaaufzv, lageso-*); keep only active workflow steps.
    ("seed_bank", "Seed Bank"),
    ("master_bank_registry", "Master Bank Registry"),
]
STEP_LABEL_LOOKUP = {step: label for step, label in ASSURED_STEP_ORDER}


def _resolve_pdf_path(path_str: str) -> Optional[Path]:
    candidate = Path(path_str)
    if candidate.exists():
        return candidate
    alt = ROOT / path_str
    if alt.exists():
        return alt
    return None


def _ordered_pdf_entries(manifest: Dict[str, object]) -> List[Dict[str, object]]:
    reports = manifest.get("files", {}).get("reports", [])
    grouped: Dict[str, List[Dict[str, object]]] = {}
    for entry in reports:
        if not isinstance(entry, dict):
            continue
        if entry.get("type") and entry.get("type") != "pdf":
            continue
        step = entry.get("step") or "misc"
        grouped.setdefault(step, []).append(entry)

    for entries in grouped.values():
        entries.sort(key=lambda item: item.get("timestamp", 0))

    ordered: List[Dict[str, object]] = []
    for step, _ in ASSURED_STEP_ORDER:
        ordered.extend(grouped.pop(step, []))

    remaining: List[Dict[str, object]] = []
    for entries in grouped.values():
        remaining.extend(entries)
    remaining.sort(key=lambda item: item.get("timestamp", 0))
    ordered.extend(remaining)
    return ordered


def _project_root(project_id: str) -> Path:
    return ensure_dirs(project_id)


def _bundle_summary(manifest: Dict[str, object], generated_at: int) -> Dict[str, object]:
    files = manifest.get("files", {})
    return {
        "project_id": manifest.get("meta", {}).get("project_id"),
        "generated_at": generated_at,
        "meta": manifest.get("meta", {}),
        "file_counts": {category: len(entries) for category, entries in files.items()},
        "chain_entries": len(manifest.get("chain", [])),
        "audit_entries": len(manifest.get("audit", [])),
    }


def generate_project_bundle(project_id: str) -> Path:
    root = _project_root(project_id)
    manifest = load_manifest(project_id)
    exports_dir = project_subdir(project_id, "exports")
    timestamp_suffix = int(time.time())
    bundle_name = f"{project_id}_bundle_{timestamp_suffix}.zip"
    bundle_path = exports_dir / bundle_name

    folders_to_include = ["snapshots", "reports", "uploads", "chainproofs"]
    with zipfile.ZipFile(bundle_path, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        for folder in folders_to_include:
            folder_path = root / folder
            if folder_path.exists():
                for file_path in folder_path.rglob("*"):
                    if file_path.is_file():
                        archive.write(file_path, arcname=str(file_path.relative_to(root)))
        manifest_file = manifest_path(project_id)
        if manifest_file.exists():
            archive.write(manifest_file, arcname="manifest.json")
        summary_payload = _bundle_summary(manifest, timestamp_suffix)
        archive.writestr("bundle_summary.json", json.dumps(summary_payload, indent=2))
    return bundle_path


def generate_project_summary_doc(project_id: str, manifest: Dict[str, object]) -> Path:
    exports_dir = project_subdir(project_id, "exports")
    timestamp_suffix = int(time.time())
    filename = f"{project_id}_summary_{timestamp_suffix}.docx"
    doc_path = exports_dir / filename

    document = Document()
    document.add_heading(f"Project Summary - {project_id}", level=0)

    meta = manifest.get("meta", {})
    document.add_heading("Meta", level=1)
    for key, value in meta.items():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{key}: ").bold = True
        paragraph.add_run(str(value))

    files = manifest.get("files", {})
    document.add_heading("Artifacts", level=1)
    if files:
        for category, entries in files.items():
            document.add_heading(category, level=2)
            table = document.add_table(rows=1, cols=4)
            headers = ["Label", "Filename", "SHA-256", "Timestamp"]
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header
            for entry in entries:
                row = table.add_row().cells
                row[0].text = str(entry.get("label", "-"))
                row[1].text = str(entry.get("filename", "-"))
                row[2].text = str(entry.get("sha256", "-"))
                row[3].text = str(entry.get("timestamp", "-"))
    else:
        document.add_paragraph("No files registered yet.")

    chain_entries = manifest.get("chain", [])
    document.add_heading("On-chain Activity", level=1)
    if chain_entries:
        for entry in chain_entries:
            paragraph = document.add_paragraph()
            paragraph.add_run("Step: ").bold = True
            paragraph.add_run(str(entry.get("step", "-")))
            paragraph.add_run(" | Tx: ").bold = True
            paragraph.add_run(str(entry.get("tx_hash", "-")))
            paragraph.add_run(" | Hash: ").bold = True
            paragraph.add_run(str(entry.get("digest", "-")))
    else:
        document.add_paragraph("No transactions recorded.")

    audit_entries = manifest.get("audit", [])
    document.add_heading("Audit Trail", level=1)
    if audit_entries:
        for entry in audit_entries:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{entry.get('timestamp', '-')}: ").bold = True
            paragraph.add_run(str(entry.get("event", "-")))
            if "file" in entry:
                paragraph.add_run(" | File: ").bold = True
                paragraph.add_run(str(entry["file"]))
            if "tx_hash" in entry:
                paragraph.add_run(" | Tx: ").bold = True
                paragraph.add_run(str(entry["tx_hash"]))
    else:
        document.add_paragraph("No audit entries logged.")

    document.save(doc_path)
    return doc_path


def _register_export(
    project_id: str,
    artifact_type: str,
    path: Path,
    digest: str,
    metadata: Optional[Dict[str, object]] = None,
) -> None:
    payload = {
        "label": artifact_type,
        "filename": path.name,
        "path": str(path),
        "sha256": digest,
        "timestamp": int(time.time()),
    }
    if metadata:
        payload.update(metadata)
    register_file(project_id, "exports", payload)
    append_audit(
        project_id,
        {
            "timestamp": int(time.time()),
            "event": f"{artifact_type.lower()}_generated",
            "file": str(path),
            "sha256": digest,
        },
    )


def _anchor_artifact(project_id: str, step: str, artifact: Dict[str, object]) -> None:
    if artifact.get("tx_hash"):
        st.info("Already anchored on-chain.")
        return
    try:
        result = send_log_tx(artifact["digest"], step, Path(artifact["path"]).as_uri())
    except Exception as exc:
        st.error(f"Failed to anchor {artifact['filename']}: {exc}")
        return

    artifact["tx_hash"] = result["tx_hash"]
    st.success(f"Anchored {artifact['filename']} (tx: {result['tx_hash']})")
    register_chain_tx(
        project_id,
        {
            "timestamp": int(time.time()),
            "step": step,
            "tx_hash": result["tx_hash"],
            "digest": artifact["digest"],
            "file": artifact["path"],
            "metadata_uri": Path(artifact["path"]).as_uri(),
        },
    )
    append_audit(
        project_id,
        {
            "timestamp": int(time.time()),
            "event": f"{step.lower()}_anchored",
            "tx_hash": result["tx_hash"],
            "file": artifact["path"],
            "digest": artifact["digest"],
        },
    )


PAGE_TITLE = "Project Dashboard"

init_page(PAGE_TITLE)
st.title(PAGE_TITLE)
st.caption("Overview of all active projects, their manifests, and close-out utilities.")
st.page_link("pages/15_Assured_Protocol_Wizard.py", label="Launch protocol wizard")

projects = list_projects()
st.metric("Total projects", len(projects))

status_counts: Dict[str, int] = {}
for project_id in projects:
    manifest = load_manifest(project_id)
    status = manifest.get("meta", {}).get("status", "unknown")
    status_counts[status] = status_counts.get(status, 0) + 1

if status_counts:
    st.subheader("Status distribution")
    st.bar_chart(status_counts)

selected_project = use_project("Project")

if not selected_project:
    st.info("Select or create a project from the sidebar to inspect details.")
    st.stop()

st.header(f"Project: {selected_project}")
manifest = load_manifest(selected_project)
col_meta, col_files, col_chain = st.columns(3)

with col_meta:
    st.subheader("Meta")
    st.json(manifest.get("meta", {}), expanded=False)

with col_files:
    st.subheader("Files")
    st.json(manifest.get("files", {}), expanded=False)

with col_chain:
    st.subheader("Chain Activity")
    st.json(manifest.get("chain", []), expanded=False)

st.subheader("Audit log")
audit_entries = manifest.get("audit", [])
if not audit_entries:
    st.caption("No audit entries yet.")
else:
    st.json(audit_entries, expanded=False)

st.subheader("Close-out toolkit")

# Auto-repair manifests: scan for orphaned PDFs
if st.button("🔧 Scan & repair manifest", help="Find PDFs in reports/ not registered in manifest"):
    sys.path.append(str(ROOT))
    from scripts.repair_manifests import repair_manifest
    result = repair_manifest(selected_project, dry_run=False)
    if result["status"] == "repaired":
        st.success(f"✅ Registered {result['added']} orphaned PDF(s)")
        for file_info in result.get("files", []):
            st.caption(f"  [{file_info['step']}] {file_info['path']}")
        st.rerun()
    elif result["status"] == "clean":
        st.info("All PDFs already registered")
    else:
        st.warning("Repair check failed")

if st.button("📚 Regenerate binders (all projects)", help="Merge latest step PDFs into binders for all projects"):
    projects = list_projects()
    results = []
    for pid in projects:
        try:
            results.append(build_binder(pid))
        except Exception as e:
            results.append({"status": "error", "project_id": pid, "error": str(e)})
    ok = sum(1 for r in results if r.get("status") == "ok")
    none = sum(1 for r in results if r.get("status") == "no-candidates")
    st.success(f"Binder regeneration complete: ok={ok}, no-candidates={none}")
    # Build quick summary table
    rows = []
    for r in results:
        pid = r.get("project_id", "-")
        status = r.get("status", "-")
        included = r.get("binder", {}).get("included_count") if status == "ok" else None
        skipped = ", ".join(r.get("binder", {}).get("skipped", [])) if status == "ok" else ""
        rows.append({"Project": pid, "Status": status, "Included": included, "Skipped": skipped})
    st.dataframe(rows, use_container_width=True)

    # Show per-project binder download buttons for successful regenerations
    st.caption("Latest binders:")
    for r in results:
        if r.get("status") == "ok":
            b = r.get("binder", {})
            path_str = b.get("path")
            if path_str:
                p = Path(path_str)
                if p.exists():
                    with p.open("rb") as fh:
                        st.download_button(
                            label=f"Download binder → {p.name}",
                            data=fh,
                            file_name=p.name,
                            key=f"download_binder_{r.get('project_id','unknown')}_{p.name}",
                        )
                    # Provide a quick link to the containing folder (local file URI)
                    st.markdown(
                        f"Open folder: [" + str(p.parent) + "](\"" + p.parent.as_uri() + "\")",
                        help="Opens the local folder if supported by your environment",
                    )
                    # Copy full path to clipboard
                    st.code(str(p), language="text")
                    st.caption("Copy path above to clipboard if needed.")

bundle_state = st.session_state.setdefault("project_bundles", {})
summary_state = st.session_state.setdefault("project_summaries", {})
bundle_info = bundle_state.get(selected_project)
summary_info = summary_state.get(selected_project)

col_bundle, col_summary = st.columns(2)

with col_bundle:
    if st.button("Generate project bundle (ZIP)", key=f"bundle_{selected_project}"):
        bundle_path = generate_project_bundle(selected_project)
        bundle_digest = sha256_bytes(bundle_path.read_bytes())
        bundle_info = {
            "filename": bundle_path.name,
            "path": str(bundle_path),
            "digest": bundle_digest,
            "timestamp": int(time.time()),
        }
        bundle_state[selected_project] = bundle_info
        _register_export(selected_project, "ProjectBundle", bundle_path, bundle_digest)
        st.success(f"Bundle created: {bundle_path.name}")

    if bundle_info:
        path = Path(bundle_info["path"])
        if path.exists():
            with path.open("rb") as handle:
                st.download_button(
                    "Download bundle",
                    handle,
                    file_name=path.name,
                    key=f"bundle_download_{selected_project}",
                )
        st.code(f"{bundle_info['filename']} SHA-256: {bundle_info['digest']}", language="text")
        if st.button("Anchor project bundle", key=f"bundle_anchor_{selected_project}"):
            _anchor_artifact(selected_project, "ProjectBundle", bundle_info)

with col_summary:
    if st.button("Generate project summary (DOCX)", key=f"summary_{selected_project}"):
        summary_path = generate_project_summary_doc(selected_project, manifest)
        summary_digest = sha256_bytes(summary_path.read_bytes())
        summary_info = {
            "filename": summary_path.name,
            "path": str(summary_path),
            "digest": summary_digest,
            "timestamp": int(time.time()),
        }
        summary_state[selected_project] = summary_info
        _register_export(selected_project, "ProjectSummary", summary_path, summary_digest)
        st.success(f"Summary report created: {summary_path.name}")

    if summary_info:
        path = Path(summary_info["path"])
        if path.exists():
            with path.open("rb") as handle:
                st.download_button(
                    "Download summary",
                    handle,
                    file_name=path.name,
                    key=f"summary_download_{selected_project}",
                )
        st.code(f"{summary_info['filename']} SHA-256: {summary_info['digest']}", language="text")
        if st.button("Anchor project summary", key=f"summary_anchor_{selected_project}"):
            _anchor_artifact(selected_project, "ProjectSummary", summary_info)

st.subheader("ASSURED binder (merged PDF)")
pdf_entries_ordered = _ordered_pdf_entries(manifest)

# De-duplicate: keep only latest PDF per step (by highest timestamp)
latest_per_step: Dict[str, Dict[str, object]] = {}
for entry in pdf_entries_ordered:
    if entry.get("step") == "assured_binder":
        continue
    resolved = _resolve_pdf_path(entry.get("path", ""))
    if not resolved:
        continue
    step_id = entry.get("step") or "misc"
    timestamp = entry.get("timestamp", 0)
    # Keep only the latest entry per step
    if step_id not in latest_per_step or timestamp > latest_per_step[step_id].get("timestamp", 0):
        label = STEP_LABEL_LOOKUP.get(step_id, step_id.replace("_", " ").replace("-", " ").title())
        latest_per_step[step_id] = {
            "step": step_id,
            "label": label,
            "path": resolved,
            "timestamp": timestamp,
        }

# Build binder candidates from de-duplicated entries, preserving step order
binder_candidates: List[Dict[str, object]] = []
for step, _ in ASSURED_STEP_ORDER:
    if step in latest_per_step:
        binder_candidates.append(latest_per_step[step])
# Add any remaining steps not in ASSURED_STEP_ORDER
for step_id, entry in latest_per_step.items():
    if not any(s == step_id for s, _ in ASSURED_STEP_ORDER):
        binder_candidates.append(entry)

if binder_candidates:
    st.caption("Included PDFs (in merge order):")
    for item in binder_candidates:
        st.write(f"• {item['label']} — {item['path'].name}")
else:
    st.info("Generate the step PDFs (Design, Delivery, Assessment, …) to enable the merged binder export.")

present_steps = {item["step"] for item in binder_candidates}
missing_steps = [label for step, label in ASSURED_STEP_ORDER if step not in present_steps]
if missing_steps:
    st.caption("Missing steps: " + ", ".join(missing_steps))

binder_entries_manifest = [
    entry
    for entry in manifest.get("files", {}).get("reports", [])
    if isinstance(entry, dict) and entry.get("step") == "assured_binder" and entry.get("type") == "pdf"
]
binder_entries_manifest.sort(key=lambda item: item.get("timestamp", 0), reverse=True)
latest_binder_entry = binder_entries_manifest[0] if binder_entries_manifest else None

if st.button(
    "Generate merged ASSURED binder",
    key=f"binder_generate_{selected_project}",
    disabled=not binder_candidates,
):
    try:
        candidate_paths = [item["path"] for item in binder_candidates]
        valid_paths: List[Path] = []
        skipped: List[Path] = []  # Paths missing or failing PDF parse
        for p in candidate_paths:
            if not p.exists():
                skipped.append(p)
                continue
            if PdfReader is not None:
                try:
                    # Attempt to parse; will raise if truncated/corrupt.
                    with p.open("rb") as fh:
                        PdfReader(fh)
                    valid_paths.append(p)
                except Exception:
                    skipped.append(p)
            else:
                valid_paths.append(p)  # cannot validate without PyPDF2
        if not valid_paths:
            st.error("No valid PDF files to merge. All candidates were missing or invalid.")
        else:
            binder_dir = project_subdir(selected_project, "reports", "binders")
            timestamp_now = int(time.time())
            binder_filename = build_step_filename(selected_project, "assured-binder", timestamp_now)
            binder_path = binder_dir / binder_filename
            merge_pdfs(valid_paths, binder_path)
            binder_digest = sha256_bytes(binder_path.read_bytes())
            binder_payload = {
                "step": "assured_binder",
                "path": str(binder_path),
                "timestamp": timestamp_now,
                "digest": binder_digest,
                "type": "pdf",
                "skipped": [p.name for p in skipped],
                "included_count": len(valid_paths),
            }
            register_file(selected_project, "reports", binder_payload)
            manifest.setdefault("files", {}).setdefault("reports", []).append(binder_payload)
            latest_binder_entry = binder_payload
            st.success(f"ASSURED binder created: {binder_filename}")
            if skipped:
                st.warning(
                    f"Skipped {len(skipped)} invalid/missing PDF(s): "
                    + ", ".join(s.name for s in skipped)
                )
    except RuntimeError as err:
        st.error(str(err))
    except ValueError:
        st.warning("No PDFs available to merge.")

if latest_binder_entry:
    binder_path = _resolve_pdf_path(latest_binder_entry.get("path", ""))
    if binder_path and binder_path.exists():
        with binder_path.open("rb") as handle:
            st.download_button(
                "Download merged binder (PDF)",
                handle,
                file_name=binder_path.name,
                key=f"assured_binder_download_{selected_project}",
            )
        st.caption(f"Merged binder saved: {binder_path}")
        skipped_list = latest_binder_entry.get("skipped") or []
        if skipped_list:
            st.caption("Skipped in last merge: " + ", ".join(skipped_list))

st.subheader("LAGESO documentation")
files_manifest = manifest.get("files", {})
report_entries = files_manifest.get("reports", [])
snapshot_entries = files_manifest.get("snapshots", [])
lageso_excel_reports = [entry for entry in report_entries if entry.get("step") == "lageso_excel"]
lageso_audit_snapshots = [entry for entry in snapshot_entries if entry.get("step") == "lageso_audit"]
lageso_chain = [
    entry for entry in manifest.get("chain", []) if entry.get("step") in {"lageso_audit", "LAGESO Audit"}
]
form_z_reports = [entry for entry in report_entries if entry.get("step") == "form_z"]
form_z_snapshots = [entry for entry in snapshot_entries if entry.get("step") == "form_z"]
form_z_chain = [entry for entry in manifest.get("chain", []) if entry.get("step") == "form_z"]

col_excel, col_audit, col_chain = st.columns(3)
if lageso_excel_reports:
    latest_excel = max(lageso_excel_reports, key=lambda e: e.get("timestamp", 0))
    col_excel.metric("Excel generated", "Yes", help=f"Latest: {latest_excel.get('filename', '')}")
else:
    col_excel.metric("Excel generated", "No")

if lageso_audit_snapshots:
    latest_audit = max(lageso_audit_snapshots, key=lambda e: e.get("timestamp", 0))
    col_audit.metric("Audit snapshot", "Yes", help=f"Snapshot: {latest_audit.get('path', '')}")
else:
    col_audit.metric("Audit snapshot", "No")

if lageso_chain:
    latest_tx = max(lageso_chain, key=lambda e: e.get("timestamp", 0))
    col_chain.metric("Anchored", "Yes", help=f"Tx: {latest_tx.get('tx_hash', '')}")
else:
    col_chain.metric("Anchored", "No")

try:
    st.page_link("pages/10_LAGESO_Compliance.py", label="Open Compliance Manager")
    st.page_link("pages/11_LAGESO_Audit_Log.py", label="Open Audit Documentation")
    st.page_link("pages/12_Form_Z_Manager.py", label="Open Form Z Manager")
except Exception:
    st.caption(
        " • ".join(
            [
                "Open Compliance Manager (via sidebar)",
                "Open Audit Documentation (via sidebar)",
                "Open Form Z Manager (via sidebar)",
            ]
        )
    )

st.subheader("Form Z documentation")
col_formz_excel, col_formz_snapshot, col_formz_chain = st.columns(3)
if form_z_reports:
    latest_formz_report = max(form_z_reports, key=lambda e: e.get("timestamp", 0))
    col_formz_excel.metric("Form Z Excel", "Yes", help=f"Latest: {latest_formz_report.get('path', '')}")
else:
    col_formz_excel.metric("Form Z Excel", "No")

if form_z_snapshots:
    latest_formz_snapshot = max(form_z_snapshots, key=lambda e: e.get("timestamp", 0))
    col_formz_snapshot.metric("Form Z snapshot", "Yes", help=f"Snapshot: {latest_formz_snapshot.get('path', '')}")
else:
    col_formz_snapshot.metric("Form Z snapshot", "No")

if form_z_chain:
    latest_formz_tx = max(form_z_chain, key=lambda e: e.get("timestamp", 0))
    col_formz_chain.metric("Form Z anchored", "Yes", help=f"Tx: {latest_formz_tx.get('tx_hash', '')}")
else:
    col_formz_chain.metric("Form Z anchored", "No")
