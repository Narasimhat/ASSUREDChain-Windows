import sys
import json
import time
import hashlib
from pathlib import Path
from typing import List, Optional

import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.components.layout import init_page

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.append(str(ROOT))

from app.components.web3_client import send_log_tx

REPORT_DIR = ROOT / "data" / "reports"
REPORT_DIR.mkdir(parents=True, exist_ok=True)
ATTACHMENTS_DIR = REPORT_DIR / "assets"
ATTACHMENTS_DIR.mkdir(parents=True, exist_ok=True)
LOGO_PATH = ROOT / "assets" / "bih_logo.png"


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def add_heading(doc: Document, text: str, level: int = 0) -> None:
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_key_value(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(value or "N/A")


def build_report(payload: dict, output_path: Path, logo: Optional[Path]) -> Path:
    meta = payload.get("meta", {})
    primers = payload.get("primers", [])
    expected = payload.get("expected_bands_bp", {})
    gel_image = payload.get("gel_image")
    lanes = payload.get("gel_lanes", [])
    sanger_traces = payload.get("sanger", [])
    call = payload.get("call", "Review")
    notes = payload.get("notes", "")

    document = Document()
    if logo and logo.exists():
        try:
            document.add_picture(str(logo), width=Inches(1.6))
        except Exception:
            pass

    add_heading(document, "Technology Platform Pluripotent Stem Cell", 0)
    add_heading(document, "Report of Intended Editing - Master Bank Verification", 1)

    add_key_value(document, "Cell line name", meta.get("cell_line", ""))
    add_key_value(document, "Bank", meta.get("bank_id", ""))
    add_key_value(document, "Passage No.", meta.get("passage", ""))
    add_key_value(document, "Genetic locus", meta.get("locus", ""))
    add_key_value(document, "Transgene / Mutation", meta.get("intended_edit", ""))
    add_key_value(document, "Name operator", meta.get("author", ""))
    add_key_value(document, "Date of testing", meta.get("timestamp", ""))

    document.add_heading("Primers", level=2)
    table = document.add_table(rows=1, cols=4)
    header = table.rows[0].cells
    header[0].text = "Name (oligo no.)"
    header[1].text = "Sequence"
    header[2].text = "Direction"
    header[3].text = "Amplicon"

    for primer in primers:
        expected_amplicon = str(
            primer.get("expected_amplicon_bp") or expected.get("WT") or ""
        )
        if primer.get("forward"):
            row = table.add_row().cells
            row[0].text = primer.get("name", "")
            row[1].text = primer.get("forward", "")
            row[2].text = "Forward"
            row[3].text = expected_amplicon
        if primer.get("reverse"):
            row = table.add_row().cells
            row[0].text = primer.get("name", "")
            row[1].text = primer.get("reverse", "")
            row[2].text = "Reverse"
            row[3].text = expected_amplicon

    document.add_heading("Results", level=2)
    document.add_paragraph("2% standard agarose gel with DNA stain; run conditions per SOP.")
    if expected:
        document.add_paragraph(
            f"Expected bands (bp): WT = {expected.get('WT', 'N/A')}, "
            f"Edited = {expected.get('edited', 'N/A')}"
        )

    if gel_image and Path(gel_image).exists():
        try:
            document.add_paragraph("PCR picture:")
            document.add_picture(str(gel_image), width=Inches(5.0))
        except Exception:
            document.add_paragraph("[Gel image omitted]")

    if lanes:
        document.add_paragraph("Lanes:")
        for lane in lanes:
            observed = lane.get("observed_bands_bp", [])
            observed_txt = ", ".join(str(o) for o in observed) if observed else "N/A"
            document.add_paragraph(
                f"- {lane.get('lane', '')}: observed {observed_txt}"
            )

    if sanger_traces:
        add_heading(document, "Sanger sequencing traces", 2)
        for idx, trace in enumerate(sanger_traces, start=1):
            document.add_paragraph(f"Trace {idx}: {trace.get('summary', '')}")
            trace_file = trace.get("file")
            if trace_file and Path(trace_file).exists() and str(trace_file).lower().endswith(
                (".png", ".jpg", ".jpeg")
            ):
                try:
                    document.add_picture(str(trace_file), width=Inches(5.0))
                except Exception:
                    document.add_paragraph("[Trace omitted]")

    add_heading(document, "Conclusion", 2)
    document.add_paragraph(
        f"The cell line {meta.get('cell_line', '')} {meta.get('bank_id', '')} shows editing verification call: "
        f"{call}. {notes or ''}"
    )
    document.add_paragraph("")
    document.add_paragraph(
        f"Responsible person: {meta.get('author', '')}              date: {meta.get('timestamp', '')}"
    )

    document.save(output_path)
    return output_path


def save_upload(upload, prefix: str) -> Optional[Path]:
    if not upload:
        return None
    extension = Path(upload.name).suffix
    filename = f"{int(time.time())}_{prefix}{extension}"
    destination = ATTACHMENTS_DIR / filename
    destination.write_bytes(upload.getvalue())
    upload.seek(0)
    return destination


def parse_observed_bands(raw: str) -> List[str]:
    values = [chunk.strip() for chunk in raw.split(",") if chunk.strip()]
    return values


init_page("Step 9 - Editing Verification (Master Bank) Report Generator")
st.title("Step 9 - Editing Verification (Master Bank) Report Generator")
st.caption(
    "Fill in the Editing Verification data below to generate a formatted DOCX report "
    "and optionally anchor it on-chain."
)

with st.form("ev_form"):
    cols = st.columns(3)
    with cols[0]:
        cell_line = st.text_input("Cell line", value="BIHi005-A-1X")
        bank_id = st.text_input("Bank ID", value="MB01")
        passage = st.text_input("Passage", value="P16")
    with cols[1]:
        project_id = st.text_input("Project ID", value="SORCS1_KI_v1")
        genetic_locus = st.text_input("Genetic locus", value="SORCS1 Exon 25")
        intended_edit = st.text_input("Intended edit", value="H1098* (KI)")
    with cols[2]:
        author = st.text_input("Responsible person", value="Narasimha Telugu")
        timestamp = st.date_input("Date of testing").strftime("%Y-%m-%d")
        call = st.selectbox("Editing verification call", ["Pass", "Review", "Fail"], index=1)

    st.markdown("#### Expected bands (bp)")
    exp_cols = st.columns(2)
    with exp_cols[0]:
        expected_wt = st.text_input("WT band", value="428")
    with exp_cols[1]:
        expected_edit = st.text_input("Edited band", value="")

    st.markdown("#### Primers")
    primer_count = st.number_input("Number of primer pairs", min_value=1, max_value=5, value=2, step=1)
    primers = []
    for idx in range(primer_count):
        with st.expander(f"Primer pair {idx + 1}", expanded=idx < 2):
            name = st.text_input(f"Name / oligo no. {idx + 1}", key=f"primer_name_{idx}", value=f"Primer {idx + 1}")
            forward_seq = st.text_input("Forward sequence", key=f"primer_forward_{idx}")
            reverse_seq = st.text_input("Reverse sequence", key=f"primer_reverse_{idx}")
            expected_amplicon = st.text_input(
                "Expected amplicon (bp)", key=f"primer_amplicon_{idx}", value=expected_wt
            )
            primers.append(
                {
                    "name": name,
                    "forward": forward_seq,
                    "reverse": reverse_seq,
                    "expected_amplicon_bp": expected_amplicon,
                }
            )

    st.markdown("#### Gel lanes")
    lane_count = st.number_input("Number of gel lanes to document", min_value=0, max_value=8, value=2, step=1)
    lanes = []
    for idx in range(lane_count):
        with st.expander(f"Lane {idx + 1}", expanded=idx < 2):
            lane_name = st.text_input("Lane label", key=f"lane_name_{idx}", value=f"Lane {idx + 1}")
            observed_raw = st.text_input(
                "Observed bands (comma-separated)", key=f"lane_observed_{idx}", value=expected_wt if idx == 0 else ""
            )
            lanes.append(
                {
                    "lane": lane_name,
                    "observed_bands_bp": parse_observed_bands(observed_raw),
                }
            )

    gel_image_upload = st.file_uploader("PCR gel image (optional)", type=["png", "jpg", "jpeg"])

    st.markdown("#### Sanger traces (optional)")
    sanger_count = st.number_input("Number of Sanger traces", min_value=0, max_value=5, value=0, step=1)
    sanger_traces = []
    for idx in range(sanger_count):
        with st.expander(f"Sanger trace {idx + 1}", expanded=False):
            summary = st.text_input("Summary", key=f"sanger_summary_{idx}")
            trace_file = st.file_uploader("Trace image", type=["png", "jpg", "jpeg"], key=f"sanger_file_{idx}")
            sanger_traces.append(
                {
                    "summary": summary,
                    "file_upload": trace_file,
                }
            )

    notes = st.text_area("Notes / comments", value="")

    submitted = st.form_submit_button("Generate report")

if submitted:
    timestamp_unix = int(time.time())
    saved_gel = save_upload(gel_image_upload, "gel") if gel_image_upload else None

    processed_sanger = []
    for idx, trace in enumerate(sanger_traces):
        file_path = save_upload(trace.get("file_upload"), f"sanger_{idx}")
        processed_sanger.append(
            {
                "summary": trace.get("summary", ""),
                "file": file_path.as_posix() if file_path else None,
            }
        )

    payload = {
        "meta": {
            "project_id": project_id,
            "cell_line": cell_line,
            "bank_id": bank_id,
            "passage": passage,
            "locus": genetic_locus,
            "intended_edit": intended_edit,
            "author": author,
            "timestamp": timestamp,
            "timestamp_unix": timestamp_unix,
        },
        "primers": primers,
        "expected_bands_bp": {"WT": expected_wt, "edited": expected_edit},
        "gel_image": saved_gel.as_posix() if saved_gel else None,
        "gel_lanes": lanes,
        "sanger": processed_sanger,
        "call": call,
        "notes": notes,
    }

    snapshot_path = REPORT_DIR / f"editver_mb_snapshot_{timestamp_unix}.json"
    snapshot_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")

    report_name = f"EditingVerification_MB_Report_{timestamp_unix}.docx"
    report_path = REPORT_DIR / report_name
    build_report(payload, report_path, LOGO_PATH if LOGO_PATH.exists() else None)

    st.success("Report generated.")
    st.caption(f"Snapshot saved: {snapshot_path}")
    st.code(f"Snapshot SHA-256: {sha256_bytes(snapshot_path.read_bytes())}", language="text")

    with report_path.open("rb") as report_bytes:
        st.download_button(
            "Download Editing Verification (MB) Report (DOCX)",
            report_bytes,
            file_name=report_name,
        )

    digest = sha256_bytes(report_path.read_bytes())
    st.code(f"Report SHA-256: {digest}", language="text")

    if st.button("Anchor report on-chain (Sepolia)"):
        result = send_log_tx(
            hex_digest=digest,
            step="Editing Verification (MB) Report",
            metadata_uri=report_path.resolve().as_uri(),
        )
        st.success("Anchored on-chain.")
        st.write(f"Tx: {result['tx_hash']}")
        st.json(result["receipt"])
