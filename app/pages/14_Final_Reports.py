import json
import sys
import time
import hashlib
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches

from app.components.layout import init_page

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.append(str(ROOT))

from app.components.project_state import (
    append_audit,
    project_subdir,
    register_chain_tx,
    register_file,
    use_project,
)
from app.components.web3_client import send_log_tx
from app.components.file_utils import snapshot_to_pdf

GLOBAL_REPORT_DIR = ROOT / "data" / "reports"
GLOBAL_REPORT_DIR.mkdir(parents=True, exist_ok=True)
GLOBAL_SNAPSHOT_DIR = GLOBAL_REPORT_DIR / "snapshots"
GLOBAL_SNAPSHOT_DIR.mkdir(parents=True, exist_ok=True)
LOGO_PATH = ROOT / "assets" / "bih_logo.png"


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def add_key_value(doc: Document, label: str, value: Optional[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(value if value not in (None, "") else "-")


def add_table(doc: Document, headers: List[str], rows: List[List[Optional[str]]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = str(value if value not in (None, "") else "-")


def build_coa_doc(meta: Dict[str, str], tests: Dict[str, str], karyotype: Dict[str, str], str_block: Dict[str, str]) -> Document:
    document = Document()
    if LOGO_PATH.exists():
        try:
            document.add_picture(str(LOGO_PATH), width=Inches(1.6))
        except Exception:
            pass

    document.add_heading("Certificate of Analysis", level=0)

    add_key_value(document, "CELL LINE NAME", meta.get("cell_line"))
    add_key_value(document, "hPSCreg Link", meta.get("hpscreg_link"))
    add_key_value(document, "ALTERNATIVE NAME", meta.get("alt_name"))
    add_key_value(document, "DONOR GENDER/AGE", meta.get("donor"))
    add_key_value(document, "DISEASE / GENETIC VARIANT", meta.get("genetic_variant"))
    add_key_value(
        document,
        "BANK",
        f"Master Bank, ID {meta.get('bank_id')} , Passage {meta.get('passage')} , Freezing Date: {meta.get('freeze_date')}",
    )
    add_key_value(document, "FREEZING METHOD", meta.get("freezing_method", "Bambanker"))
    add_key_value(
        document,
        "CULTURE PLATFORM",
        f"Feeder independent | Medium: {meta.get('medium', 'E8')} | Coating: {meta.get('coating', 'Geltrex')}",
    )
    add_key_value(document, "GENETIC MODIFICATION", "yes" if meta.get("edited") else "no")
    document.add_paragraph("")

    document.add_heading("Test Description", level=1)
    headers = ["TEST", "Method", "Specification", "Result"]
    rows = [
        [
            "STERILITY (mycoplasma)",
            tests.get("myco_method", "RT-PCR"),
            "No contamination detected",
            tests.get("myco_result", "Pass"),
        ],
        [
            "STERILITY (bacteria/yeast/fungi)",
            tests.get("bact_method", "Culture 4 days, antibiotic-free"),
            "No contamination detected",
            tests.get("bact_result", "Pass"),
        ],
        [
            "VIABILITY / MORPHOLOGY",
            "Phase-contrast 24/48/72 h",
            "Typical hPSC growth",
            tests.get("morph_result", "Pass"),
        ],
        [
            "UNDIFFERENTIATED PHENOTYPE",
            tests.get("pluri_method", "IF/FACS (>=3 markers)"),
            "Markers detected",
            tests.get("pluri_result", "not done"),
        ],
        [
            "DIFFERENTIATION POTENTIAL (3 GL)",
            tests.get("tri_method", "Spontaneous / Directed"),
            "Markers of 3 germ layers",
            tests.get("tri_result", "not done"),
        ],
    ]
    add_table(document, headers, rows)
    document.add_paragraph("")

    document.add_heading("KARYOTYPE", level=1)
    add_key_value(document, "Platform", karyotype.get("platform", "SNP array"))
    add_key_value(document, "Result", karyotype.get("summary", "No major structural aberration detected"))
    add_key_value(document, "Comparison", karyotype.get("vs_parental", "Karyotype matches Parental Cell Line"))
    document.add_paragraph("")

    document.add_heading("IDENTITY (STR ANALYSIS)", level=1)
    add_key_value(document, "Method", str_block.get("method", "Promega GenePrint 10"))
    add_key_value(document, "Result", str_block.get("result", "Pass"))
    add_key_value(document, "Comparison to primary", str_block.get("match_primary", "Identical"))
    document.add_paragraph("")

    add_key_value(document, "REFERENCE", meta.get("reference_url"))
    add_key_value(document, "Date", meta.get("report_date"))
    add_key_value(document, "Signature", meta.get("signatory"))
    return document


def build_cts_doc(cts: Dict[str, str]) -> Document:
    document = Document()
    if LOGO_PATH.exists():
        try:
            document.add_picture(str(LOGO_PATH), width=Inches(1.6))
        except Exception:
            pass

    document.add_heading("Cell / Tissue Transfer Form", level=0)

    add_key_value(document, "SAMPLE NAME", cts.get("sample_name"))
    add_key_value(document, "SAMPLE TYPE", cts.get("sample_type", "iPSCs"))
    add_key_value(document, "Derived from", cts.get("derived_from", "Fibroblasts"))
    add_key_value(document, "DONOR GENDER/AGE", cts.get("donor"))
    add_key_value(document, "INFORMED CONSENT", cts.get("consent", "Yes"))
    add_key_value(document, "BSL", cts.get("bsl", "1"))
    add_key_value(document, "DISEASE / MODIFICATION", cts.get("disease", ""))
    add_key_value(document, "CELL LINE INITIALLY OBTAINED FROM", cts.get("parental"))
    add_key_value(document, "PASSAGE / CELL NUMBER", cts.get("passage_cellnum"))
    add_key_value(document, "CULTURE VESSEL COATING", cts.get("coating", "Geltrex"))
    add_key_value(document, "CULTURE MEDIUM", cts.get("medium", "E8"))
    add_key_value(document, "PASSAGING METHOD/RATE", cts.get("passaging", "PBS/EDTA 0.5 mM, 4-5 min, 1:6-1:12"))
    add_key_value(
        document,
        "THAWING INSTRUCTIONS",
        cts.get("thawing", "Seed 1 vial into <=2 wells of a 6-well plate with 10 uM ROCKi"),
    )
    add_key_value(document, "PROVIDER", cts.get("provider", "Max Delbruck Center - Stem Cell Core"))
    add_key_value(document, "CONTACT", cts.get("contact", ""))
    add_key_value(document, "DATE / SIGNATURE", cts.get("date_sig", ""))
    return document


def _safe_stem(value: str) -> str:
    cleaned = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in (value or ""))
    cleaned = cleaned.strip("_")
    return cleaned or "report"


def _write_snapshot(directory: Path, stem: str, payload: Dict[str, object]) -> Path:
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / f"{stem}.json"
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return path


def _state_bucket(key: str, project_key: str) -> List[Dict[str, object]]:
    store = st.session_state.setdefault(key, {})
    if project_key not in store:
        store[project_key] = []
    return store[project_key]


def _register_outputs(
    project_id: Optional[str],
    step: str,
    report_path: Path,
    report_digest: str,
    snapshot_path: Path,
    snapshot_digest: str,
    context: Dict[str, str],
    extras: Optional[List[Dict[str, Any]]] = None,
) -> None:
    timestamp_now = int(time.time())
    if not project_id:
        return
    register_file(
        project_id,
        "reports",
        {
            "label": step,
            "filename": report_path.name,
            "path": str(report_path),
            "sha256": report_digest,
            "timestamp": timestamp_now,
            "format": report_path.suffix.lstrip("."),
            **context,
        },
    )
    register_file(
        project_id,
        "snapshots",
        {
            "label": step,
            "filename": snapshot_path.name,
            "path": str(snapshot_path),
            "sha256": snapshot_digest,
            "timestamp": timestamp_now,
            "source_report": report_path.name,
            **context,
        },
    )
    if extras:
        for extra in extras:
            register_file(
                project_id,
                extra.get("category", "reports"),
                {
                    "label": step,
                    "timestamp": timestamp_now,
                    **extra.get("payload", {}),
                },
            )
    append_audit(
        project_id,
        {
            "timestamp": timestamp_now,
            "event": f"{step.lower()}_generated",
            "file": str(report_path),
            "snapshot": str(snapshot_path),
            "sha256": report_digest,
            "context": context,
        },
    )


def _render_generated_outputs(
    entries: List[Dict[str, object]],
    step: str,
    project_id: Optional[str],
    button_key: str,
) -> None:
    if not entries:
        return

    st.markdown("**Downloads**")
    bundle_stamp = entries[0].get("generated_at", int(time.time()))

    if len(entries) > 1:
        import io
        import zipfile

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
            for entry in entries:
                path = Path(entry["path"])
                if path.exists():
                    archive.write(path, arcname=path.name)
                pdf_info = entry.get("pdf")
                if pdf_info:
                    pdf_path = Path(pdf_info["path"])
                    if pdf_path.exists():
                        archive.write(pdf_path, arcname=pdf_path.name)
        zip_buffer.seek(0)
        st.download_button(
            f"Download all {step} documents (ZIP)",
            zip_buffer,
            file_name=f"{step.lower()}_reports_{bundle_stamp}.zip",
            key=f"{button_key}_zip",
        )

    for entry in entries:
        path = Path(entry["path"])
        if not path.exists():
            st.warning(f"{entry['filename']} could not be located on disk.")
            continue
        with path.open("rb") as handle:
            st.download_button(
                f"Download {entry['filename']}",
                handle,
                file_name=entry["filename"],
                key=f"{button_key}_download_{entry['filename']}",
            )
        st.code(f"{entry['filename']} SHA-256: {entry['digest']}", language="text")
        pdf_info = entry.get("pdf")
        if pdf_info:
            pdf_path = Path(pdf_info["path"])
            if pdf_path.exists():
                with pdf_path.open("rb") as handle:
                    st.download_button(
                        f"Download {pdf_info['filename']}",
                        handle,
                        file_name=pdf_info["filename"],
                        key=f"{button_key}_download_{pdf_info['filename']}",
                    )
                st.code(f"{pdf_info['filename']} SHA-256: {pdf_info['digest']}", language="text")
        st.caption(f"Snapshot: {entry['snapshot']}")
        if entry.get("tx_hash"):
            st.caption(f"Anchored transaction: {entry['tx_hash']}")

    if st.button(f"Anchor {step} documents on-chain", key=f"{button_key}_anchor"):
        anchored: List[Dict[str, str]] = []
        for entry in entries:
            if entry.get("tx_hash"):
                anchored.append({"file": entry["filename"], "tx_hash": str(entry["tx_hash"])})
                continue
            try:
                result = send_log_tx(entry["digest"], step, Path(entry["path"]).as_uri())
            except Exception as exc:
                st.error(f"Failed to anchor {entry['filename']}: {exc}")
                continue
            entry["tx_hash"] = result["tx_hash"]
            anchored.append({"file": entry["filename"], "tx_hash": result["tx_hash"]})
            if project_id:
                register_chain_tx(
                    project_id,
                    {
                        "timestamp": int(time.time()),
                        "step": step,
                        "tx_hash": result["tx_hash"],
                        "digest": entry["digest"],
                        "file": entry["path"],
                        "metadata_uri": Path(entry["path"]).as_uri(),
                    },
                )
                append_audit(
                    project_id,
                    {
                        "timestamp": int(time.time()),
                        "event": f"{step.lower()}_anchored",
                        "tx_hash": result["tx_hash"],
                        "file": entry["path"],
                        "digest": entry["digest"],
                    },
                )
        if anchored:
            st.success(f"Anchored {step} documents on-chain.")
            st.json(anchored)


init_page("Step 14 - Final Reports (CoA & CTS)")
st.title("Step 14 - Final Reports (CoA & CTS)")
st.caption("Assemble CoA and CTS reports, export DOCX, and anchor on-chain when ready.")

project_id = use_project("Project")
project_key = project_id or "__global__"

if project_id:
    report_dir = project_subdir(project_id, "reports", "final_reports")
    snapshot_dir = project_subdir(project_id, "snapshots", "final_reports")
else:
    report_dir = GLOBAL_REPORT_DIR
    snapshot_dir = GLOBAL_SNAPSHOT_DIR

coa_entries = _state_bucket("coa_reports", project_key)

with st.expander("1) Build Certificate of Analysis"):
    base_meta: Dict[str, str] = {}
    col1, col2, col3 = st.columns(3)
    with col1:
        base_meta["cell_line"] = st.text_input("Cell line", "BIHi005-A-99", key=f"coa_cell_line_{project_key}")
    with col2:
        base_meta["bank_id"] = st.text_input("Bank ID", "MB01", key=f"coa_bank_id_{project_key}")
    with col3:
        base_meta["passage"] = st.text_input("Passage", "29", key=f"coa_passage_{project_key}")

    base_meta["freeze_date"] = st.text_input(
        "Freezing Date (YYYY-MM-DD)", "2023-10-30", key=f"coa_freeze_date_{project_key}"
    )
    base_meta["hpscreg_link"] = st.text_input(
        "hPSCreg Link", "https://hpscreg.eu/cell-line/BIHi005-A-99", key=f"coa_hpscreg_link_{project_key}"
    )
    base_meta["alt_name"] = st.text_input(
        "Alternative name", "BIHi005-A TSC2_AA Cl.34", key=f"coa_alt_name_{project_key}"
    )
    base_meta["donor"] = st.text_input("Donor gender/age", "Male, 25-29", key=f"coa_donor_{project_key}")
    base_meta["genetic_variant"] = st.text_input(
        "Disease / genetic variant", "1364, 1365 SS to AA, heart defects", key=f"coa_variant_{project_key}"
    )
    base_meta["freezing_method"] = st.text_input(
        "Freezing method", "Bambanker", key=f"coa_freezing_method_{project_key}"
    )
    base_meta["medium"] = st.text_input("Medium", "E8", key=f"coa_medium_{project_key}")
    base_meta["coating"] = st.text_input("Coating", "Geltrex", key=f"coa_coating_{project_key}")
    base_meta["edited"] = st.checkbox("Genetic modification", True, key=f"coa_edited_{project_key}")
    base_meta["reference_url"] = st.text_input(
        "Reference URL", "https://hpscreg.eu/cell-line/BIHi005-A", key=f"coa_reference_url_{project_key}"
    )
    base_meta["report_date"] = st.text_input(
        "Report date", time.strftime("%Y-%m-%d"), key=f"coa_report_date_{project_key}"
    )
    base_meta["signatory"] = st.text_input("Signature (name)", "M. Wendt", key=f"coa_signatory_{project_key}")

    bank_id_input = st.text_input(
        "Bank IDs (comma-separated)", "MB01, MB02", key=f"coa_bank_ids_{project_key}"
    )
    bank_ids = [item.strip() for item in bank_id_input.split(",") if item.strip()]

    cell_line_input = st.text_input(
        "Cell line names (comma-separated)", base_meta["cell_line"], key=f"coa_cell_line_list_{project_key}"
    )
    cell_line_names = [item.strip() for item in cell_line_input.split(",") if item.strip()]
    if not bank_ids:
        bank_ids = [base_meta["bank_id"]]
    if not cell_line_names:
        cell_line_names = [base_meta["cell_line"]]

    preview_count = max(len(bank_ids), len(cell_line_names), 1)
    preview_rows = []
    for idx in range(preview_count):
        preview_rows.append(
            {
                "Cell line": cell_line_names[idx] if idx < len(cell_line_names) else "",
                "Bank ID": bank_ids[idx] if idx < len(bank_ids) else "",
                "hPSCreg Link": base_meta.get("hpscreg_link", ""),
                "Alternative name": base_meta.get("alt_name", ""),
            }
        )
    pairs_df = pd.DataFrame(preview_rows)
    st.markdown("**CoA pairs (edit as needed)**")
    edited_pairs_df = st.data_editor(
        pairs_df,
        hide_index=True,
        num_rows="dynamic",
        column_config={
            "Cell line": st.column_config.TextColumn(required=True),
            "Bank ID": st.column_config.TextColumn(required=True),
            "hPSCreg Link": st.column_config.TextColumn(required=False),
            "Alternative name": st.column_config.TextColumn(required=False),
        },
        key=f"coa_pairs_editor_{project_key}",
    )
    coa_pairs = [
        {
            "cell_line": str(row.get("Cell line", "")).strip(),
            "bank_id": str(row.get("Bank ID", "")).strip(),
            "hpscreg_link": str(row.get("hPSCreg Link", "") or "").strip(),
            "alt_name": str(row.get("Alternative name", "") or "").strip(),
        }
        for row in edited_pairs_df.to_dict("records")
        if row.get("Cell line") and row.get("Bank ID")
    ]
    if not coa_pairs:
        st.info("Provide at least one Cell line / Bank ID pair below.")

    st.markdown("**Test results**")
    tests = {
        "myco_method": st.text_input(
            "Mycoplasma method", "IDEXX BioAnalytics RealTime-PCR", key=f"coa_myco_method_{project_key}"
        ),
        "myco_result": st.selectbox(
            "Mycoplasma", ["Pass", "Fail", "not done"], 0, key=f"coa_myco_result_{project_key}"
        ),
        "bact_method": st.text_input(
            "Bacteria/Yeast/Fungi method",
            "4 days culture in antibiotic-free medium",
            key=f"coa_bact_method_{project_key}",
        ),
        "bact_result": st.selectbox(
            "Bacteria/Yeast/Fungi", ["Pass", "Fail", "not done"], 0, key=f"coa_bact_result_{project_key}"
        ),
        "morph_result": st.selectbox(
            "Morphology/Viability", ["Pass", "Fail", "not done"], 0, key=f"coa_morph_result_{project_key}"
        ),
        "pluri_method": st.text_input(
            "Undifferentiated phenotype method", "IF / FACS", key=f"coa_pluri_method_{project_key}"
        ),
        "pluri_result": st.selectbox(
            "Undifferentiated phenotype", ["Pass", "Fail", "not done"], 2, key=f"coa_pluri_result_{project_key}"
        ),
        "tri_method": st.text_input(
            "3-GL method", "Spontaneous / Directed", key=f"coa_tri_method_{project_key}"
        ),
        "tri_result": st.selectbox(
            "3-GL", ["Pass", "Fail", "not done"], 2, key=f"coa_tri_result_{project_key}"
        ),
    }

    st.markdown("**Karyotype & STR**")
    karyotype = {
        "platform": st.text_input(
            "Karyotype platform", "Illumina Infinium GSA-24", key=f"coa_karyo_platform_{project_key}"
        ),
        "summary": st.text_input(
            "Karyotype result",
            "No major structural aberration detected (compared to Parental line)",
            key=f"coa_karyo_summary_{project_key}",
        ),
        "vs_parental": st.text_input(
            "Comparison", "Karyotype matches Parental Cell Line", key=f"coa_karyo_vs_parental_{project_key}"
        ),
    }
    str_block = {
        "method": st.text_input("STR method", "Promega GenePrint 10", key=f"coa_str_method_{project_key}"),
        "result": st.selectbox(
            "STR analyzed", ["Pass", "Fail", "not done"], 0, key=f"coa_str_result_{project_key}"
        ),
        "match_primary": st.text_input(
            "Match to primary", "Identical to profile of primary cells", key=f"coa_str_match_{project_key}"
        ),
    }

    if st.button("Generate CoA DOCX", key=f"generate_coa_{project_key}"):
        if not coa_pairs:
            st.warning("Add at least one valid Cell line / Bank ID pair before generating.")
            st.stop()
        timestamp_suffix = int(time.time())
        new_entries: List[Dict[str, object]] = []

        for idx, pair in enumerate(coa_pairs, start=1):
            bank_id = pair["bank_id"]
            cell_line = pair["cell_line"]
            meta = base_meta.copy()
            meta["bank_id"] = bank_id
            meta["cell_line"] = cell_line
            if pair.get("hpscreg_link"):
                meta["hpscreg_link"] = pair["hpscreg_link"]
            if pair.get("alt_name"):
                meta["alt_name"] = pair["alt_name"]

            document = build_coa_doc(meta, tests, karyotype, str_block)
            safe_cell = _safe_stem(cell_line)
            safe_bank = _safe_stem(bank_id)
            stem = f"CoA_{safe_cell}_{safe_bank}_{timestamp_suffix}_{idx + 1}"
            filename = f"{stem}.docx"
            output_path = report_dir / filename
            document.save(output_path)
            digest = sha256_bytes(output_path.read_bytes())

            snapshot_payload = {
                "type": "CoA",
                "meta": meta,
                "tests": tests.copy(),
                "karyotype": karyotype.copy(),
                "str_analysis": str_block.copy(),
                "generated_at": timestamp_suffix,
                "filename": filename,
                "sha256": digest,
            }
            snapshot_path = _write_snapshot(snapshot_dir, stem, snapshot_payload)
            snapshot_digest = sha256_bytes(snapshot_path.read_bytes())

            pdf_filename = f"{stem}.pdf"
            pdf_path = report_dir / pdf_filename
            snapshot_to_pdf(
                snapshot_payload,
                pdf_path,
                "Certificate of Analysis",
            )
            pdf_digest = sha256_bytes(pdf_path.read_bytes())

            entry = {
                "filename": filename,
                "path": str(output_path),
                "digest": digest,
                "snapshot": str(snapshot_path),
                "generated_at": timestamp_suffix,
                "meta": meta.copy(),
                "pdf": {
                    "filename": pdf_filename,
                    "path": str(pdf_path),
                    "digest": pdf_digest,
                },
            }
            new_entries.append(entry)

            _register_outputs(
                project_id,
                "CoA",
                output_path,
                digest,
                snapshot_path,
                snapshot_digest,
                {"bank_id": bank_id, "cell_line": cell_line},
                extras=[
                    {
                        "category": "reports",
                        "payload": {
                            "filename": pdf_filename,
                            "path": str(pdf_path),
                            "sha256": pdf_digest,
                            "format": "pdf",
                            "bank_id": bank_id,
                            "cell_line": cell_line,
                        },
                    }
                ],
            )

        coa_entries.clear()
        coa_entries.extend(new_entries)
        st.success(f"Generated {len(new_entries)} Certificate of Analysis document(s).")

    _render_generated_outputs(coa_entries, "CoA", project_id, f"coa_{project_key}")


cts_entries = _state_bucket("cts_reports", project_key)

with st.expander("2) Build Cell / Tissue Transfer Form"):
    base_cts: Dict[str, str] = {
        "sample_name": st.text_input("Sample name", "BIHi005-A-99", key=f"cts_sample_name_{project_key}"),
        "sample_type": st.text_input("Sample type", "iPSCs", key=f"cts_sample_type_{project_key}"),
        "derived_from": st.text_input("Derived from", "Fibroblasts", key=f"cts_derived_from_{project_key}"),
        "donor": st.text_input("Donor gender/age", "Male, 25-29", key=f"cts_donor_{project_key}"),
        "consent": st.text_input("Informed consent", "Yes", key=f"cts_consent_{project_key}"),
        "bsl": st.text_input("BSL", "1", key=f"cts_bsl_{project_key}"),
        "disease": st.text_input("Disease / modification", "TSC2_AA Cl.34", key=f"cts_disease_{project_key}"),
        "parental": st.text_input("Parental line", "BIHi005-A", key=f"cts_parental_{project_key}"),
        "passage_cellnum": st.text_input(
            "Passage / Cell number", "29 / 0.968 x 10^6 ml^-1", key=f"cts_passage_cellnum_{project_key}"
        ),
        "coating": st.text_input("Coating", "Geltrex", key=f"cts_coating_{project_key}"),
        "medium": st.text_input("Medium", "E8 Home", key=f"cts_medium_{project_key}"),
        "passaging": st.text_input(
            "Passaging", "PBS/EDTA 0.5 mM, 4-5 min, 1:6-1:12", key=f"cts_passaging_{project_key}"
        ),
        "thawing": st.text_input(
            "Thawing instructions",
            "Seed 1 vial into <=2 wells of a 6-well plate with 10 uM ROCKi",
            key=f"cts_thawing_{project_key}",
        ),
        "provider": st.text_input(
            "Provider", "Max Delbruck Center - Stem Cell Core Facility", key=f"cts_provider_{project_key}"
        ),
        "contact": st.text_input(
            "Contact email/phone",
            "Sebastian.diecke@mdc-berlin.de / +49 30 9406 3090",
            key=f"cts_contact_{project_key}",
        ),
        "date_sig": st.text_input(
            "Date / Signature", f"{time.strftime('%Y-%m-%d')} / Carolin Genehr", key=f"cts_date_sig_{project_key}"
        ),
    }

    sample_names_input = st.text_input(
        "Sample names (comma-separated)", base_cts["sample_name"], key=f"cts_sample_names_{project_key}"
    )
    sample_names = [item.strip() for item in sample_names_input.split(",") if item.strip()]

    identifier_input = st.text_input(
        "CTS identifiers (comma-separated, e.g., MB01, MB02)",
        "MB01, MB02",
        key=f"cts_identifiers_{project_key}",
    )
    identifiers = [item.strip() for item in identifier_input.split(",") if item.strip()]

    if not sample_names:
        sample_names = [base_cts["sample_name"]]
    if not identifiers:
        identifiers = ["MB01"]

    cts_preview_count = max(len(sample_names), len(identifiers), 1)
    cts_preview_rows = [
        {
            "Sample name": sample_names[idx] if idx < len(sample_names) else "",
            "Identifier": identifiers[idx] if idx < len(identifiers) else "",
        }
        for idx in range(cts_preview_count)
    ]
    st.markdown("**CTS pairs (edit as needed)**")
    cts_pairs_df = st.data_editor(
        pd.DataFrame(cts_preview_rows),
        hide_index=True,
        num_rows="dynamic",
        column_config={
            "Sample name": st.column_config.TextColumn(required=True),
            "Identifier": st.column_config.TextColumn(required=True),
        },
        key=f"cts_pairs_editor_{project_key}",
    )
    cts_pairs = [
        {"sample_name": row["Sample name"].strip(), "identifier": row["Identifier"].strip()}
        for row in cts_pairs_df.to_dict("records")
        if row.get("Sample name") and row.get("Identifier")
    ]
    if not cts_pairs:
        st.info("Provide at least one Sample name / Identifier pair below.")

    if st.button("Generate CTS DOCX", key=f"generate_cts_{project_key}"):
        if not cts_pairs:
            st.warning("Add at least one valid Sample name / Identifier pair before generating.")
            st.stop()
        timestamp_suffix = int(time.time())
        new_entries: List[Dict[str, object]] = []

        for idx, pair in enumerate(cts_pairs, start=1):
            sample_name = pair["sample_name"]
            identifier = pair["identifier"]

            cts_meta = base_cts.copy()
            cts_meta["sample_name"] = sample_name
            cts_meta["identifier"] = identifier

            document = build_cts_doc(cts_meta)
            safe_sample = _safe_stem(sample_name)
            safe_identifier = _safe_stem(identifier)
            stem = f"CTS_{safe_sample}_{safe_identifier}_{timestamp_suffix}_{idx + 1}"
            filename = f"{stem}.docx"
            output_path = report_dir / filename
            document.save(output_path)
            digest = sha256_bytes(output_path.read_bytes())

            snapshot_payload = {
                "type": "CTS",
                "meta": cts_meta,
                "generated_at": timestamp_suffix,
                "filename": filename,
                "sha256": digest,
            }
            snapshot_path = _write_snapshot(snapshot_dir, stem, snapshot_payload)
            snapshot_digest = sha256_bytes(snapshot_path.read_bytes())

            pdf_filename = f"{stem}.pdf"
            pdf_path = report_dir / pdf_filename
            snapshot_to_pdf(
                snapshot_payload,
                pdf_path,
                "Cell / Tissue Transfer Form",
            )
            pdf_digest = sha256_bytes(pdf_path.read_bytes())

            entry = {
                "filename": filename,
                "path": str(output_path),
                "digest": digest,
                "snapshot": str(snapshot_path),
                "generated_at": timestamp_suffix,
                "meta": cts_meta.copy(),
                "pdf": {
                    "filename": pdf_filename,
                    "path": str(pdf_path),
                    "digest": pdf_digest,
                },
            }
            new_entries.append(entry)

            _register_outputs(
                project_id,
                "CTS",
                output_path,
                digest,
                snapshot_path,
                snapshot_digest,
                {"identifier": identifier, "sample_name": sample_name},
                extras=[
                    {
                        "category": "reports",
                        "payload": {
                            "filename": pdf_filename,
                            "path": str(pdf_path),
                            "sha256": pdf_digest,
                            "format": "pdf",
                            "identifier": identifier,
                            "sample_name": sample_name,
                        },
                    }
                ],
            )

        cts_entries.clear()
        cts_entries.extend(new_entries)
        st.success(f"Generated {len(new_entries)} CTS document(s).")

    _render_generated_outputs(cts_entries, "CTS", project_id, f"cts_{project_key}")
